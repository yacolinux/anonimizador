import logging
import os
import re
import json
import uuid
import hmac
import time
import unicodedata
import subprocess
import tempfile
import threading
import shlex
from urllib.request import Request, urlopen
from urllib.error import URLError
from io import BytesIO
from functools import wraps
from flask import session, g
from flask_session import Session
import redis
import pdf2image
import pytesseract

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s: %(message)s'
)
logger = logging.getLogger('anonimizador')

import pdfplumber
import docx
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from fpdf import FPDF
from fpdf.errors import FPDFException
from flask import (
    Flask, render_template, request, jsonify,
    send_from_directory, send_file
)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = '/app/uploads'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
app.config['SECRET_KEY'] = os.environ.get('FLASK_SECRET_KEY')
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = os.environ.get('SESSION_COOKIE_SECURE', '0') == '1'

MODEL_NAME = os.environ.get('MODEL_NAME', 'qwen/qwen3-30b-a3b')
ADMIN_USER = os.environ.get('ADMIN_USER', 'adminanon')
ADMIN_PASS = os.environ.get('ADMIN_PASS')
READY_MAX_INFLIGHT = int(os.environ.get('READY_MAX_INFLIGHT', '2'))
UPLOAD_TTL_SECONDS = int(os.environ.get('UPLOAD_TTL_SECONDS', str(900)))
LOGIN_MAX_ATTEMPTS = int(os.environ.get('LOGIN_MAX_ATTEMPTS', '5'))
LOGIN_WINDOW_SECONDS = int(os.environ.get('LOGIN_WINDOW_SECONDS', '300'))
SESSION_BACKEND = os.environ.get('SESSION_BACKEND', 'redis').lower()
REDIS_URL = os.environ.get('REDIS_URL', 'redis://redis:6379/0')
REDIS_CONFIG_KEY = os.environ.get('REDIS_CONFIG_KEY', 'anonimizador:config')
USE_AYMURAI = os.environ.get('USE_AYMURAI', '0') == '1'
AYMURAI_BASE_URL = os.environ.get('AYMURAI_BASE_URL', 'http://aymurai:8899').strip().rstrip('/')
AYMURAI_TIMEOUT_SECONDS = int(os.environ.get('AYMURAI_TIMEOUT_SECONDS', '20'))
AYMURAI_MIN_SEGMENT_CHARS = int(os.environ.get('AYMURAI_MIN_SEGMENT_CHARS', '15'))

_api_call_log = []
_MAX_API_LOG_ENTRIES = 100
_API_LOGS_REDIS_KEY = 'anonimizador:api_logs'


def _read_api_logs_from_redis():
    if not redis_available:
        return []
    try:
        entries = redis_client.lrange(_API_LOGS_REDIS_KEY, 0, -1)
        logs = []
        for raw in entries:
            if isinstance(raw, bytes):
                raw = raw.decode('utf-8')
            try:
                logs.append(json.loads(raw))
            except json.JSONDecodeError:
                continue
        return logs
    except redis.RedisError:
        return []

def _log_api_call(model_url, model_name, status, duration_ms, response_preview=''):
    entry = {
        'timestamp': time.strftime('%Y-%m-%dT%H:%M:%S'),
        'model_url': model_url,
        'model_name': model_name,
        'status': status,
        'duration_ms': duration_ms,
        'response_preview': (response_preview or '')[:200],
    }
    _api_call_log.append(entry)
    while len(_api_call_log) > _MAX_API_LOG_ENTRIES:
        _api_call_log.pop(0)
    if redis_available:
        try:
            pipe = redis_client.pipeline()
            pipe.rpush(_API_LOGS_REDIS_KEY, json.dumps(entry, ensure_ascii=False))
            pipe.ltrim(_API_LOGS_REDIS_KEY, -_MAX_API_LOG_ENTRIES, -1)
            pipe.execute()
        except redis.RedisError:
            logger.warning('No se pudo persistir el log de API en Redis')
LOCAL_INFERENCE_MAX = int(os.environ.get('LOCAL_INFERENCE_MAX', '3'))
LOCAL_INFERENCE_WAIT_SECONDS = int(os.environ.get('LOCAL_INFERENCE_WAIT_SECONDS', '90'))
LOCAL_INFERENCE_POLL_SECONDS = float(os.environ.get('LOCAL_INFERENCE_POLL_SECONDS', '1.5'))
LOCAL_INFERENCE_SLOT_TTL_SECONDS = int(os.environ.get('LOCAL_INFERENCE_SLOT_TTL_SECONDS', '180'))
MAX_UPLOAD_MB = int(os.environ.get('MAX_UPLOAD_MB', '100'))
OCR_MAX_PAGES = int(os.environ.get('OCR_MAX_PAGES', '50'))
OCR_DPI = int(os.environ.get('OCR_DPI', '200'))
OCR_LANG = os.environ.get('OCR_LANG', 'spa')

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

app.config['MAX_CONTENT_LENGTH'] = MAX_UPLOAD_MB * 1024 * 1024

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

REGEX_PATTERNS_FILE = '/app/regex_patterns.json'

_inflight_lock = threading.Lock()
_inflight_requests = 0
_login_attempts = {}
redis_client = None
redis_available = False
session_ext = Session()


def init_redis_client():
    global redis_client, redis_available
    try:
        client = redis.Redis.from_url(
            REDIS_URL,
            decode_responses=False,
            socket_connect_timeout=2,
            socket_timeout=2,
        )
        client.ping()
        redis_client = client
        redis_available = True
        logger.info('Redis conectado en %s', REDIS_URL)
    except redis.RedisError as e:
        redis_available = False
        redis_client = None
        logger.warning('Redis no disponible (%s). Se usa fallback local.', e)


def configure_session_backend():
    if SESSION_BACKEND == 'redis' and redis_available:
        app.config['SESSION_TYPE'] = 'redis'
        app.config['SESSION_REDIS'] = redis_client
        app.config['SESSION_PERMANENT'] = False
        app.config['SESSION_USE_SIGNER'] = True
        session_ext.init_app(app)
        logger.info('Sesiones configuradas en Redis')
    else:
        logger.info('Sesiones en cookie (backend local)')


def bootstrap_config_in_redis():
    if not redis_available:
        return
    try:
        existing = redis_client.get(REDIS_CONFIG_KEY)
        if existing:
            return
        try:
            with open(REGEX_PATTERNS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            data = DEFAULT_PATTERNS_DATA
        redis_client.set(REDIS_CONFIG_KEY, json.dumps(data, ensure_ascii=False))
        logger.info('Config inicial cargada en Redis')
    except redis.RedisError:
        logger.warning('No se pudo bootstrapear config en Redis')


def validate_required_env():
    required = []
    if not app.config.get('SECRET_KEY'):
        required.append('FLASK_SECRET_KEY')
    if not ADMIN_PASS:
        required.append('ADMIN_PASS')
    if required:
        raise RuntimeError(
            f"Missing required environment variables: {', '.join(required)}"
        )


def cleanup_old_uploads():
    cutoff = time.time() - UPLOAD_TTL_SECONDS
    upload_dir = app.config['UPLOAD_FOLDER']
    try:
        for name in os.listdir(upload_dir):
            path = os.path.join(upload_dir, name)
            if not os.path.isfile(path):
                continue
            if os.path.getmtime(path) < cutoff:
                os.unlink(path)
    except OSError:
        logger.warning('No se pudo ejecutar limpieza de uploads antiguos')


def is_valid_upload_filename(filename):
    return bool(re.fullmatch(r'[0-9a-fA-F-]{36}\.(pdf|docx)', filename or ''))


def is_path_inside_uploads(path):
    upload_root = os.path.realpath(app.config['UPLOAD_FOLDER'])
    target = os.path.realpath(path)
    return target.startswith(upload_root + os.sep)


def _prune_login_attempts(now):
    stale = [ip for ip, ts in _login_attempts.items() if now - ts[-1] > LOGIN_WINDOW_SECONDS]
    for ip in stale:
        _login_attempts.pop(ip, None)


def is_login_rate_limited(ip):
    if redis_available:
        key = f'anonimizador:login:{ip}'
        try:
            count = redis_client.get(key)
            return int(count or 0) >= LOGIN_MAX_ATTEMPTS
        except redis.RedisError:
            pass
    now = time.time()
    _prune_login_attempts(now)
    attempts = _login_attempts.get(ip, [])
    attempts = [t for t in attempts if now - t <= LOGIN_WINDOW_SECONDS]
    _login_attempts[ip] = attempts
    return len(attempts) >= LOGIN_MAX_ATTEMPTS


def register_login_failure(ip):
    if redis_available:
        key = f'anonimizador:login:{ip}'
        try:
            pipe = redis_client.pipeline()
            pipe.incr(key)
            pipe.expire(key, LOGIN_WINDOW_SECONDS)
            pipe.execute()
            return
        except redis.RedisError:
            pass
    now = time.time()
    attempts = _login_attempts.get(ip, [])
    attempts.append(now)
    _login_attempts[ip] = attempts


validate_required_env()
init_redis_client()
configure_session_backend()
bootstrap_config_in_redis()


def _cleanup_loop():
    while True:
        time.sleep(60)
        try:
            cleanup_old_uploads()
        except Exception:
            logger.warning('Error en limpieza periódica de uploads')


cleanup_thread = threading.Thread(target=_cleanup_loop, daemon=True, name='cleanup-loop')
cleanup_thread.start()


def _inc_inflight():
    global _inflight_requests
    with _inflight_lock:
        _inflight_requests += 1
        return _inflight_requests


def _dec_inflight():
    global _inflight_requests
    with _inflight_lock:
        _inflight_requests = max(0, _inflight_requests - 1)
        return _inflight_requests


def _current_inflight():
    with _inflight_lock:
        return _inflight_requests


@app.before_request
def track_request_start():
    if request.path == '/ready':
        g._counted_inflight = False
        return
    g._counted_inflight = True
    _inc_inflight()


@app.after_request
def track_request_end(response):
    if getattr(g, '_counted_inflight', False):
        _dec_inflight()
        g._counted_inflight = False
    return response


@app.teardown_request
def track_request_teardown(_exc):
    if getattr(g, '_counted_inflight', False):
        _dec_inflight()
        g._counted_inflight = False

DEFAULT_PATTERNS_DATA = {
    "patterns": [
        {"pattern": r"\b\d{1,2}\.?\d{3}\.?\d{3}\b", "type": "dni_argentino"},
        {"pattern": r"\b\d{7,8}\b", "type": "dni_argentino"},
        {"pattern": r"(?:calle|av\.|avenida|domicilio|direcci[oó]n|pasaje|b[oó]u?le?vard|ruta|camino)\s+[a-záéíóúñ\s]+\d+", "type": "direccion"},
        {"pattern": r"\d+\s*(?:años|anios|años de edad)", "type": "edad"},
        {"pattern": r"\b(?:masculino|femenino|var[oó]n|mujer|hombre|femenina|masculina)\b", "type": "sexo"},
        {"pattern": r"(?:paciente|nombre|apellido|señor|señora|sr[a]?\.?)\s*:?\s*[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)+", "type": "nombre"},
        {"pattern": r"\b(?:[\w.-]+@[\w.-]+\.\w{2,})\b", "type": "email"},
        {"pattern": r"\b(?:abus\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:viol\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:fallec\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:homicid\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:femicid\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:lesion\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:amenaz\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:agred\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:imput\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:conden\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:deten\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:testig\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:denunci\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:perici\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:forens\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:cadav\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:autops\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:necrops\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:identif\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:document\w*)\b", "type": "sensible"},
        {"pattern": r"\b(?:expedient\w*)\b", "type": "sensible"},
    ],
    "prompt": (
        "Analiza este documento y determina TODAS las palabras o conjuntos de palabras "
        "que son datos personales o información judicial/jurídica importante.\n\n"
        "Busca ESPECÍFICAMENTE:\n"
        "- Nombres y apellidos completos de personas\n"
        "- DNI, CUIT, CUIL, pasaporte, cualquier número de identificación\n"
        "- Domicilios o direcciones postales completas\n"
        "- Edades (números seguidos de 'años')\n"
        "- Sexo o género de la persona\n"
        "- Teléfonos, celulares, emails\n"
        "- Cuentas bancarias, CBUs, números de tarjeta\n"
        "- Montos de dinero: cifras en números Y TAMBIÉN cifras escritas en palabras\n"
        "  (ej: 'cuarenta y cuatro millones de pesos', 'dos mil quinientos dólares')\n"
        "- Cualquier cifra o valor descrito en palabras o números que sea relevante\n"
        "- Fechas de nacimiento, fechas importantes\n"
        "- Nombres de empresas, juzgados, organismos\n"
        "- Números de expediente, causas judiciales\n\n"
        "IMPORTANTE: al final de tu respuesta, incluí un bloque JSON con el "
        "siguiente formato, y SOLO ese bloque JSON al final:\n\n"
        '```json\n'
        '[{"word": "el texto exacto como aparece", "type": "categoria"}, ...]\n'
        '```\n\n'
        "TEXT:\n\"\"\"\n{text}\n\"\"\""
    )
}


def load_regex_config():
    if redis_available:
        try:
            raw = redis_client.get(REDIS_CONFIG_KEY)
            if raw:
                if isinstance(raw, bytes):
                    raw = raw.decode('utf-8')
                data = json.loads(raw)
                logger.info('Config regex cargada desde Redis')
                return data
        except (redis.RedisError, json.JSONDecodeError):
            logger.warning('No se pudo leer config desde Redis, usando fallback')
    try:
        with open(REGEX_PATTERNS_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        logger.info('Config regex cargada desde archivo')
        return data
    except (FileNotFoundError, json.JSONDecodeError):
        logger.info('Usando config regex por defecto')
        return DEFAULT_PATTERNS_DATA


def get_pii_patterns():
    config = load_regex_config()
    return [(p['pattern'], p['type']) for p in config.get('patterns', [])]


def get_opencode_prompt():
    config = load_regex_config()
    return config.get('prompt', DEFAULT_PATTERNS_DATA['prompt'])


def save_regex_config(
    patterns,
    prompt,
    model_url=None,
    model_name=None,
    opencode_command=None,
    api_key=None,
    use_direct_api=None,
    use_opencode=None,
    use_aymurai=None,
    aymurai_url=None,
):
    data = {"patterns": patterns, "prompt": prompt}
    if model_url is not None:
        data["model_url"] = model_url
    if model_name is not None:
        data["model_name"] = model_name
    if opencode_command is not None:
        data["opencode_command"] = opencode_command
    elif "opencode_command" in load_regex_config():
        data["opencode_command"] = load_regex_config().get("opencode_command")
    if api_key is not None:
        data["api_key"] = api_key
    elif "api_key" in load_regex_config():
        data["api_key"] = load_regex_config().get("api_key")
    if use_direct_api is not None:
        data["use_direct_api"] = bool(use_direct_api)
    elif "use_direct_api" in load_regex_config():
        data["use_direct_api"] = load_regex_config().get("use_direct_api")
    if use_opencode is not None:
        data["use_opencode"] = bool(use_opencode)
    elif "use_opencode" in load_regex_config():
        data["use_opencode"] = load_regex_config().get("use_opencode")
    if use_aymurai is not None:
        data["use_aymurai"] = bool(use_aymurai)
    elif "use_aymurai" in load_regex_config():
        data["use_aymurai"] = load_regex_config().get("use_aymurai")
    if aymurai_url is not None:
        data["aymurai_url"] = aymurai_url
    elif "aymurai_url" in load_regex_config():
        data["aymurai_url"] = load_regex_config().get("aymurai_url")
    if redis_available:
        try:
            redis_client.set(REDIS_CONFIG_KEY, json.dumps(data, ensure_ascii=False))
            logger.info('Config regex guardada en Redis (%d patrones)', len(patterns))
        except redis.RedisError:
            logger.warning('No se pudo guardar config en Redis, guardando solo en archivo')
    with open(REGEX_PATTERNS_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    logger.info('Config regex guardada (%d patrones)', len(patterns))


def get_model_config():
    config = load_regex_config()
    return {
        'model_url': config.get('model_url', os.environ.get('OPENAI_BASE_URL', '')),
        'model_name': config.get('model_name', os.environ.get('MODEL_NAME', '')),
        'opencode_command': config.get(
            'opencode_command',
            'opencode run "{message}" --model opencode/{model} --dangerously-skip-permissions --file {file}'
        ),
        'api_key': config.get('api_key', os.environ.get('OPENAI_API_KEY', '')),
        'use_direct_api': config.get('use_direct_api', False),
        'use_opencode': config.get('use_opencode', True),
        'use_aymurai': config.get('use_aymurai', USE_AYMURAI),
    }


def get_opencode_command_template():
    cfg = get_model_config()
    return cfg['opencode_command']


def get_current_model():
    cfg = get_model_config()
    return cfg['model_name'] or MODEL_NAME


def get_current_base_url():
    cfg = get_model_config()
    return cfg['model_url'] or os.environ.get('OPENAI_BASE_URL', '')


def get_current_api_key():
    cfg = get_model_config()
    return cfg['api_key'] or os.environ.get('OPENAI_API_KEY', '')


def is_local_model_provider():
    model_url = (get_current_base_url() or '').strip().lower()
    if not model_url:
        return False
    local_hosts = (
        'localhost',
        '127.0.0.1',
        '0.0.0.0',
        'host.docker.internal',
        'ollama',
    )
    return any(host in model_url for host in local_hosts)


def local_provider_healthcheck_url():
    base = (get_current_base_url() or '').strip()
    if not base:
        return None
    if '/api/' in base:
        return base.rstrip('/')
    return base.rstrip('/') + '/api/tags'


def is_local_provider_available():
    if not is_local_model_provider():
        return True
    target = local_provider_healthcheck_url()
    if not target:
        return False
    try:
        req = Request(target, method='GET')
        with urlopen(req, timeout=3) as resp:
            return resp.status == 200
    except (URLError, TimeoutError, ValueError):
        return False


def acquire_local_inference_slot(wait_seconds=None):
    if not redis_available or not is_local_model_provider() or LOCAL_INFERENCE_MAX <= 0:
        return None, None

    token = str(uuid.uuid4())
    wait_limit = LOCAL_INFERENCE_WAIT_SECONDS if wait_seconds is None else wait_seconds
    deadline = time.time() + max(0, wait_limit)
    slot_prefix = 'anonimizador:local_inference:slot'

    while time.time() <= deadline:
        for slot_num in range(1, LOCAL_INFERENCE_MAX + 1):
            slot_key = f'{slot_prefix}:{slot_num}'
            try:
                ok = redis_client.set(
                    slot_key,
                    token,
                    nx=True,
                    ex=max(30, LOCAL_INFERENCE_SLOT_TTL_SECONDS),
                )
            except redis.RedisError:
                logger.warning('Redis error tomando slot de inferencia local')
                return None, None
            if ok:
                return slot_key, token
        time.sleep(max(0.2, LOCAL_INFERENCE_POLL_SECONDS))

    return None, token


def release_local_inference_slot(slot_key, token):
    if not redis_available or not slot_key or not token:
        return
    try:
        val = redis_client.get(slot_key)
        if val and val.decode('utf-8') == token:
            redis_client.delete(slot_key)
    except redis.RedisError:
        logger.warning('Redis error liberando slot de inferencia local')


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_ocr(filepath):
    segments = []
    try:
        logger.info('Convirtiendo PDF a imagenes para OCR (dpi=%d, lang=%s)', OCR_DPI, OCR_LANG)
        images = pdf2image.convert_from_path(filepath, dpi=OCR_DPI)
        for i, image in enumerate(images):
            if i >= OCR_MAX_PAGES:
                logger.warning('Limite de paginas OCR alcanzado (%d)', OCR_MAX_PAGES)
                break
            text = pytesseract.image_to_string(image, lang=OCR_LANG)
            if text.strip():
                segments.append({'type': 'paragraph', 'text': text.strip()})
    except Exception as e:
        logger.error('Error en OCR: %s', e)
    return segments


def extract_text_pdf(filepath):
    segments = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                try:
                    lines = page.extract_text_lines() or []
                    for block in lines:
                        raw = block.get('text', '').strip()
                        if not raw:
                            continue
                        x0 = block.get('x0', 0)
                        top = block.get('top', 0)
                        is_short = len(raw) < 90
                        high_on_page = isinstance(top, (int, float)) and top < 60
                        indented = isinstance(x0, (int, float)) and (x0 > 50)
                        if is_short and (high_on_page or indented):
                            seg_type = 'title'
                        else:
                            seg_type = 'paragraph'
                        segments.append({'type': seg_type, 'text': raw})
                except Exception:
                    raw = page.extract_text() or ''
                    for line in raw.split('\n'):
                        line = line.strip()
                        if line:
                            segments.append({'type': 'paragraph', 'text': line})
    except Exception:
        pass
    if not segments:
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    raw = page.extract_text() or ''
                    for line in raw.split('\n'):
                        line = line.strip()
                        if line:
                            segments.append({'type': 'paragraph', 'text': line})
        except Exception:
            pass
    
    total_text_len = sum(len(s['text']) for s in segments)
    if total_text_len < 100:
        logger.info('PDF con poco texto extraible (%d chars). Iniciando fallback OCR.', total_text_len)
        ocr_segments = extract_text_ocr(filepath)
        if ocr_segments:
            return ocr_segments, True
    return segments, False


def extract_text_docx(filepath):
    doc = docx.Document(filepath)
    segments = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name.lower() if para.style else ''
        if 'heading' in style_name or 'title' in style_name:
            seg_type = 'title'
        elif 'list' in style_name:
            seg_type = 'list'
        else:
            seg_type = 'paragraph'
        segments.append({'type': seg_type, 'text': text})
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                segments.append({'type': 'paragraph', 'text': ' | '.join(row_texts)})
    return segments


def extract_text(filepath):
    ext = filepath.rsplit('.', 1)[1].lower()
    if ext == 'pdf':
        return extract_text_pdf(filepath)
    elif ext == 'docx':
        return extract_text_docx(filepath), False
    return [], False


def segments_to_plaintext(segments):
    return '\n\n'.join(s['text'] for s in segments)


def get_model_id():
    model_name = get_current_model()
    parts = model_name.replace('\\', '/').split('/')
    return parts[-1] if len(parts) > 1 else parts[0]


def parse_pii_from_output(raw_output):
    text = re.sub(r'\x1b\[[0-9;]*m', '', raw_output or '').strip()

    candidates = []
    fenced = re.search(r'```(?:json)?\s*(\[.*?\])\s*```', text, re.DOTALL)
    if fenced:
        candidates.append(fenced.group(1))

    inline_array = re.search(r'(\[\s*\{.*\}\s*\])', text, re.DOTALL)
    if inline_array:
        candidates.append(inline_array.group(1))

    first_bracket = text.find('[')
    last_bracket = text.rfind(']')
    if first_bracket != -1 and last_bracket != -1 and last_bracket > first_bracket:
        candidates.append(text[first_bracket:last_bracket + 1])

    for candidate in candidates:
        try:
            parsed = json.loads(candidate)
            if isinstance(parsed, list):
                cleaned = []
                for item in parsed:
                    if isinstance(item, dict) and 'word' in item and 'type' in item:
                        word = str(item.get('word', '')).strip()
                        ptype = str(item.get('type', 'other')).strip() or 'other'
                        if word:
                            cleaned.append({'word': word, 'type': ptype})
                if cleaned:
                    return cleaned
        except json.JSONDecodeError:
            continue

    pair_matches = re.findall(
        r'\{\s*["\']word["\']\s*:\s*["\']([^"\']+)["\']\s*,\s*["\']type["\']\s*:\s*["\']([^"\']+)["\']\s*\}',
        text,
        re.IGNORECASE,
    )
    if pair_matches:
        return [{'word': w.strip(), 'type': t.strip() or 'other'} for w, t in pair_matches if w.strip()]

    table_lines = [ln.strip() for ln in text.splitlines() if '|' in ln]
    extracted = []
    for ln in table_lines:
        if ln.startswith('|---'):
            continue
        cols = [c.strip(' `') for c in ln.split('|') if c.strip()]
        if len(cols) >= 2 and cols[0].lower() not in ('palabra', 'word'):
            extracted.append({'word': cols[0], 'type': cols[1] or 'other'})
    if extracted:
        return extracted

    return []


def clean_opencode_inference_output(raw_output):
    text = re.sub(r'\x1b\[[0-9;]*m', '', raw_output or '').strip()
    if not text:
        return ''

    noise_markers = (
        'Performing one time database migration',
        'sqlite-migration:done',
        'Database migration complete.',
    )
    noise_line_prefixes = (
        '> build ',
        '> build·',
        '> build',
    )

    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            if lines and lines[-1] != '':
                lines.append('')
            continue
        if any(marker in stripped for marker in noise_markers):
            break
        if stripped.startswith(noise_line_prefixes):
            break
        lines.append(stripped)

    cleaned = '\n'.join(lines).strip()
    return cleaned or text


def call_opencode_for_pii(text, wait_seconds=None):
    model_id = get_model_id()
    prompt_template = get_opencode_prompt()
    prompt = prompt_template.replace('{text}', text)

    combined_input = (
        "INSTRUCCIONES:\n" + prompt + "\n\n"
        "NO HAGAS PREGUNTAS. NO SUGIERAS ACCIONES. Solo devolvé el JSON con "
        "los datos detectados. No escribas nada más que el JSON."
    )

    logger.info('Enviando texto a opencode (%d chars)', len(text))

    with tempfile.NamedTemporaryFile(
        mode='w', suffix='.txt', delete=False, dir='/app/uploads'
    ) as f:
        f.write(combined_input)
        tmp_path = f.name

    short_msg = 'Analiza el archivo adjunto y seguí las instrucciones'
    command_template = (get_opencode_command_template() or '').strip()
    if not command_template:
        command_template = 'opencode run "{message}" --model opencode/{model} --dangerously-skip-permissions --file {file}'
    command_filled = (
        command_template
        .replace('{message}', short_msg)
        .replace('{model}', model_id)
        .replace('{file}', tmp_path)
    )
    queue_notice = None
    slot_key = None
    slot_token = None

    if is_local_model_provider() and not is_local_provider_available():
        queue_notice = 'Proveedor local no disponible en este momento.'
        logger.warning(queue_notice)
        return [], queue_notice, queue_notice, 'unavailable'

    if is_local_model_provider() and redis_available and LOCAL_INFERENCE_MAX > 0:
        logger.info(
            'Proveedor local detectado. Intentando tomar slot de inferencia (max=%d)',
            LOCAL_INFERENCE_MAX,
        )
        started_wait = time.time()
        slot_key, slot_token = acquire_local_inference_slot(wait_seconds=wait_seconds)
        waited = int(time.time() - started_wait)
        if not slot_key:
            queue_notice = (
                'Proveedor local ocupado. No se obtuvo slot de inferencia dentro del tiempo de espera.'
            )
            logger.warning(queue_notice)
            return [], queue_notice, queue_notice, 'busy'
        if waited > 0:
            queue_notice = f'Inferencia local en cola: espera aproximada {waited}s antes de procesar.'
            logger.info(queue_notice)

    try:
        logger.info('Ejecutando comando opencode personalizado')
        cmd_args = shlex.split(command_filled)
        run_env = {**os.environ, 'HOME': os.environ.get('HOME', '/root')}
        selected_api_key = (get_current_api_key() or '').strip()
        if selected_api_key:
            run_env['OPENAI_API_KEY'] = selected_api_key
        result = subprocess.run(
            cmd_args,
            capture_output=True,
            text=True,
            timeout=120,
            env=run_env,
        )

        full_output = (result.stdout or '') + '\n' + (result.stderr or '')
        full_output = full_output.strip()
        logger.info('Respuesta opencode: %d chars', len(full_output))

        parsed = parse_pii_from_output(full_output)
        if parsed:
            logger.info('PII detectadas por IA: %d keywords', len(parsed))
            return parsed, full_output, queue_notice, 'ok'

        logger.warning('No se pudo parsear JSON del output de opencode. Preview: %s', full_output[:220])
        return [], full_output, queue_notice, 'ok'
    except subprocess.TimeoutExpired:
        logger.error('Timeout al ejecutar opencode (120s)')
        return [], 'Error: Timeout al ejecutar opencode', queue_notice, 'timeout'
    except (json.JSONDecodeError, subprocess.CalledProcessError) as e:
        logger.error('Error ejecutando opencode: %s', e)
        return [], f'Error: {str(e)}', queue_notice, 'error'
    finally:
        release_local_inference_slot(slot_key, slot_token)
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def call_direct_api_for_pii(text, wait_seconds=None):
    prompt_template = get_opencode_prompt()
    prompt = prompt_template.replace('{text}', text)

    combined_input = (
        "INSTRUCCIONES:\n" + prompt + "\n\n"
        "NO HAGAS PREGUNTAS. NO SUGIERAS ACCIONES. Solo devolvé el JSON con "
        "los datos detectados. No escribas nada más que el JSON."
    )

    cfg = get_model_config()
    base_url = (cfg['model_url'] or '').strip().rstrip('/')
    model_name = cfg.get('model_name', get_current_model())
    api_key = (cfg.get('api_key') or '').strip()

    if not base_url:
        logger.warning('Direct API: sin model_url configurado')
        return [], 'Error: no model_url configured', None, 'error'

    queue_notice = None
    slot_key = None
    slot_token = None

    if is_local_model_provider() and not is_local_provider_available():
        queue_notice = 'Proveedor local no disponible en este momento.'
        logger.warning(queue_notice)
        return [], queue_notice, queue_notice, 'unavailable'

    if is_local_model_provider() and redis_available and LOCAL_INFERENCE_MAX > 0:
        logger.info(
            'Direct API local. Intentando tomar slot de inferencia (max=%d)',
            LOCAL_INFERENCE_MAX,
        )
        started_wait = time.time()
        slot_key, slot_token = acquire_local_inference_slot(wait_seconds=wait_seconds)
        waited = int(time.time() - started_wait)
        if not slot_key:
            queue_notice = (
                'Proveedor local ocupado. No se obtuvo slot dentro del tiempo de espera.'
            )
            logger.warning(queue_notice)
            return [], queue_notice, queue_notice, 'busy'
        if waited > 0:
            queue_notice = f'Inferencia local en cola: espera aproximada {waited}s.'
            logger.info(queue_notice)

    api_url = f"{base_url}/chat/completions"
    payload = json.dumps({
        "model": model_name,
        "messages": [
            {
                "role": "system",
                "content": (
                    "Sos un asistente que detecta datos personales en documentos. "
                    "Respondé ÚNICAMENTE con el JSON solicitado."
                ),
            },
            {"role": "user", "content": combined_input},
        ],
        "temperature": 0,
        "max_tokens": 4096,
    }).encode('utf-8')

    headers = {'Content-Type': 'application/json'}
    if api_key:
        headers['Authorization'] = f'Bearer {api_key}'

    logger.info('Direct API: llamada a %s con modelo %s', api_url, model_name)
    started = time.time()
    try:
        req = Request(api_url, data=payload, headers=headers, method='POST')
        with urlopen(req, timeout=120) as resp:
            duration_ms = int((time.time() - started) * 1000)
            raw = resp.read()
            body = raw.decode('utf-8', errors='replace')
            data = json.loads(body)
            content = data.get('choices', [{}])[0].get('message', {}).get('content', '')
            logger.info('Direct API: respuesta en %dms, %d chars', duration_ms, len(content))

            _log_api_call(base_url, model_name, 'ok', duration_ms, content[:200])

            parsed = parse_pii_from_output(content)
            if parsed:
                logger.info('Direct API detectó %d keywords', len(parsed))
                return parsed, content, queue_notice, 'ok'

            logger.warning('Direct API no pudo parsear JSON. Preview: %s', content[:220])
            return [], content, queue_notice, 'ok'
    except URLError as e:
        duration_ms = int((time.time() - started) * 1000)
        _log_api_call(base_url, model_name, f'error: {e}', duration_ms)
        logger.error('Direct API URLError: %s', e)
        return [], f'Error de conexion: {e}', queue_notice, 'error'
    except subprocess.TimeoutExpired:
        duration_ms = int((time.time() - started) * 1000)
        _log_api_call(base_url, model_name, 'timeout', duration_ms)
        logger.error('Direct API timeout (120s)')
        return [], 'Error: Timeout al llamar API', queue_notice, 'timeout'
    except Exception as e:
        duration_ms = int((time.time() - started) * 1000)
        _log_api_call(base_url, model_name, f'error: {e}', duration_ms)
        logger.error('Direct API error: %s', e)
        return [], f'Error: {e}', queue_notice, 'error'
    finally:
        release_local_inference_slot(slot_key, slot_token)


def call_opencode_for_inference(prompt_text, wait_seconds=None):
    model_id = get_model_id()
    logger.info('Enviando prompt de inferencia a opencode (%d chars)', len(prompt_text))

    with tempfile.NamedTemporaryFile(
        mode='w', suffix='.txt', delete=False, dir='/app/uploads'
    ) as f:
        f.write(prompt_text)
        tmp_path = f.name

    short_msg = 'Respondé la consulta del archivo adjunto de forma breve y clara'
    command_template = (get_opencode_command_template() or '').strip()
    if not command_template:
        command_template = 'opencode run "{message}" --model opencode/{model} --dangerously-skip-permissions --file {file}'
    command_filled = (
        command_template
        .replace('{message}', short_msg)
        .replace('{model}', model_id)
        .replace('{file}', tmp_path)
    )

    slot_key = None
    slot_token = None
    if is_local_model_provider() and not is_local_provider_available():
        logger.warning('Proveedor local no disponible en este momento.')
        return None, 'Proveedor local no disponible en este momento.', 'unavailable'

    if is_local_model_provider() and redis_available and LOCAL_INFERENCE_MAX > 0:
        logger.info(
            'Inferencia opencode local. Intentando tomar slot (max=%d)',
            LOCAL_INFERENCE_MAX,
        )
        slot_key, slot_token = acquire_local_inference_slot(wait_seconds=wait_seconds)
        if not slot_key:
            logger.warning('Proveedor local ocupado. No se obtuvo slot de inferencia.')
            return None, 'Proveedor local ocupado. No se obtuvo slot de inferencia.', 'busy'

    try:
        logger.info('Ejecutando comando opencode de inferencia')
        cmd_args = shlex.split(command_filled)
        run_env = {**os.environ, 'HOME': os.environ.get('HOME', '/root')}
        selected_api_key = (get_current_api_key() or '').strip()
        if selected_api_key:
            run_env['OPENAI_API_KEY'] = selected_api_key
        result = subprocess.run(
            cmd_args,
            capture_output=True,
            text=True,
            timeout=120,
            env=run_env,
        )
        output = (result.stdout or '') + '\n' + (result.stderr or '')
        output = output.strip()
        output = clean_opencode_inference_output(output)
        logger.info('Respuesta inferencia opencode: %d chars', len(output))
        return output, None, 'ok'
    except subprocess.TimeoutExpired:
        logger.error('Timeout al ejecutar opencode de inferencia (120s)')
        return None, 'Error: Timeout al ejecutar opencode', 'timeout'
    except (json.JSONDecodeError, subprocess.CalledProcessError) as e:
        logger.error('Error ejecutando opencode de inferencia: %s', e)
        return None, f'Error: {str(e)}', 'error'
    finally:
        release_local_inference_slot(slot_key, slot_token)
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def call_direct_api_for_inference(prompt_text, wait_seconds=None):
    cfg = get_model_config()
    base_url = (cfg['model_url'] or '').strip().rstrip('/')
    model_name = cfg.get('model_name', get_current_model())
    api_key = (cfg.get('api_key') or '').strip()

    if not base_url:
        return None, 'Error: no model_url configured', 'error'

    if is_local_model_provider() and not is_local_provider_available():
        return None, 'Proveedor local no disponible en este momento.', 'unavailable'

    slot_key = None
    slot_token = None
    if is_local_model_provider() and redis_available and LOCAL_INFERENCE_MAX > 0:
        slot_key, slot_token = acquire_local_inference_slot(wait_seconds=wait_seconds)
        if not slot_key:
            return None, 'Proveedor local ocupado. No se obtuvo slot de inferencia.', 'busy'

    api_url = f"{base_url}/chat/completions"
    payload = json.dumps({
        "model": model_name,
        "messages": [
            {
                "role": "system",
                "content": "Sos un asistente útil. Respondé en español de forma breve y clara."
            },
            {"role": "user", "content": prompt_text},
        ],
        "temperature": 0.2,
        "max_tokens": 512,
    }).encode('utf-8')

    headers = {'Content-Type': 'application/json'}
    if api_key:
        headers['Authorization'] = f'Bearer {api_key}'

    started = time.time()
    try:
        req = Request(api_url, data=payload, headers=headers, method='POST')
        with urlopen(req, timeout=120) as resp:
            duration_ms = int((time.time() - started) * 1000)
            raw = resp.read()
            body = raw.decode('utf-8', errors='replace')
            data = json.loads(body)
            content = data.get('choices', [{}])[0].get('message', {}).get('content', '')
            _log_api_call(base_url, model_name, 'inference-ok', duration_ms, content[:200])
            return content, None, 'ok'
    except URLError as e:
        duration_ms = int((time.time() - started) * 1000)
        _log_api_call(base_url, model_name, f'inference-error: {e}', duration_ms)
        logger.error('Direct API inference URLError: %s', e)
        return None, f'Error de conexion: {e}', 'error'
    except subprocess.TimeoutExpired:
        duration_ms = int((time.time() - started) * 1000)
        _log_api_call(base_url, model_name, 'inference-timeout', duration_ms)
        logger.error('Direct API inference timeout (120s)')
        return None, 'Error: Timeout al llamar API', 'timeout'
    except Exception as e:
        duration_ms = int((time.time() - started) * 1000)
        _log_api_call(base_url, model_name, f'inference-error: {e}', duration_ms)
        logger.error('Direct API inference error: %s', e)
        return None, f'Error: {e}', 'error'
    finally:
        release_local_inference_slot(slot_key, slot_token)


def run_model_inference(prompt_text, wait_seconds=None):
    cfg = get_model_config()
    if cfg.get('use_direct_api'):
        return call_direct_api_for_inference(prompt_text, wait_seconds=wait_seconds)
    return call_opencode_for_inference(prompt_text, wait_seconds=wait_seconds)


def get_aymurai_url():
    config = load_regex_config()
    return (config.get('aymurai_url', '') or '').strip().rstrip('/') or AYMURAI_BASE_URL


def use_aymurai():
    if not bool(get_aymurai_url()):
        return False
    config = load_regex_config()
    return config.get('use_aymurai', USE_AYMURAI)


def map_aymurai_label_to_type(label_name):
    label_map = {
        'NOMBRE': 'nombre',
        'PER': 'nombre',
        'GENERO': 'sexo',
        'FECHA_DE_NACIMIENTO': 'fecha',
        'FECHA_RESOLUCION': 'fecha',
        'FECHA_DEL_HECHO': 'fecha',
        'HORA_DE_INICIO': 'fecha',
        'HORA_DE_CIERRE': 'fecha',
        'LUGAR_DEL_HECHO': 'direccion',
        'LOC': 'direccion',
        'DIRECCION': 'direccion',
        'DOMICILIO': 'direccion',
        'EDAD': 'edad',
        'EDAD_AL_MOMENTO_DEL_HECHO': 'edad',
        'DNI': 'dni_argentino',
        'N_EXPTE_EJE': 'sensible',
        'TIPO_DE_RESOLUCION': 'sensible',
        'OBJETO_DE_LA_RESOLUCION': 'sensible',
        'CONDUCTA': 'sensible',
        'CONDUCTA_DESCRIPCION': 'sensible',
        'ART_INFRINGIDO': 'sensible',
        'DETALLE': 'sensible',
        'FRASES_AGRESION': 'sensible',
        'VIOLENCIA_DE_GENERO': 'sensible',
        'MODALIDAD_DE_LA_VIOLENCIA': 'sensible',
        'RELACION_Y_TIPO_ENTRE_ACUSADO/A_Y_DENUNCIANTE': 'sensible',
        'PERSONA_ACUSADA_NO_DETERMINADA': 'sensible',
        'HIJOS_HIJAS_EN_COMUN': 'sensible',
        'NACIONALIDAD': 'other',
        'NIVEL_INSTRUCCION': 'other',
    }
    return label_map.get((label_name or '').strip().upper(), 'other')


def resolve_aymurai_range(text, label_text, hint_start=None):
    ranges = find_normalized_ranges(text, label_text)
    if not ranges:
        return None
    if hint_start is None:
        return ranges[0]
    return min(ranges, key=lambda item: abs(item[0] - hint_start))


def extract_aymurai_label_payload(label):
    attrs = label.get('attrs') or {}
    word = (attrs.get('aymurai_alt_text') or label.get('text') or '').strip()
    start = attrs.get('aymurai_alt_start_char')
    end = attrs.get('aymurai_alt_end_char')
    if start is None or end is None:
        start = label.get('start_char')
        end = label.get('end_char')
    return {
        'word': word,
        'start': start,
        'end': end,
        'type': map_aymurai_label_to_type(attrs.get('aymurai_label')),
        'label_name': attrs.get('aymurai_label'),
    }


def call_aymurai_for_segments(segments):
    if not use_aymurai():
        return [], []

    api_url = f'{get_aymurai_url()}/anonymizer/predict'
    keywords = []
    positions = []
    seen_keywords = set()

    for seg_idx, seg in enumerate(segments):
        text = (seg.get('text') or '').strip()
        if len(text) < AYMURAI_MIN_SEGMENT_CHARS:
            continue

        payload = json.dumps({'text': text}).encode('utf-8')
        req = Request(
            api_url,
            data=payload,
            headers={'Content-Type': 'application/json'},
            method='POST',
        )
        try:
            with urlopen(req, timeout=AYMURAI_TIMEOUT_SECONDS) as resp:
                body = resp.read().decode('utf-8', errors='replace')
                data = json.loads(body)
        except Exception as e:
            logger.warning('AymurAI no disponible para segmento %d: %s', seg_idx, e)
            return [], []

        for label in data.get('labels') or []:
            parsed = extract_aymurai_label_payload(label)
            word = parsed['word']
            if not word:
                continue

            resolved = resolve_aymurai_range(text, word, parsed['start'])
            if not resolved:
                continue
            start, end = resolved
            position = {
                'segment': seg_idx,
                'start': start,
                'end': end,
                'word': text[start:end],
                'type': parsed['type'],
            }
            positions.append(position)

            key = (normalize_text(position['word']).lower(), position['type'])
            if key not in seen_keywords:
                seen_keywords.add(key)
                keywords.append({'word': position['word'], 'type': position['type']})

    logger.info('AymurAI detectó %d keywords / %d posiciones', len(keywords), len(positions))
    return keywords, positions


def _segment_coverage_ratio(seg_text, positions_in_seg):
    if not positions_in_seg or not seg_text:
        return 0.0
    covered = set()
    for p in positions_in_seg:
        for i in range(p['start'], p['end']):
            covered.add(i)
    return len(covered) / max(len(seg_text), 1)


def _get_uncovered_segments(segments, positions, coverage_threshold=0.30):
    uncovered = []
    covered_count = 0
    for seg_idx, seg in enumerate(segments):
        seg_positions = [p for p in positions if p['segment'] == seg_idx]
        ratio = _segment_coverage_ratio(seg['text'], seg_positions)
        if ratio < coverage_threshold:
            uncovered.append(seg)
        else:
            covered_count += 1
    if covered_count > 0:
        logger.info(
            'Optimizacion IA: %d segmentos ya cubiertos (ratio>=%.0f%%), enviando %d a IA',
            covered_count, coverage_threshold * 100, len(uncovered),
        )
    return uncovered


def run_detection_pipeline(segments, wait_seconds=None):
    default_keywords, default_positions = detect_default_pii(segments)
    logger.info('Regex detecto %d keywords / %d posiciones', len(default_keywords), len(default_positions))

    aymurai_keywords, aymurai_positions = call_aymurai_for_segments(segments)

    pre_positions = default_positions + aymurai_positions
    uncovered_segments = _get_uncovered_segments(segments, pre_positions)
    plaintext = segments_to_plaintext(uncovered_segments)
    logger.info('Texto plano para IA: %d chars (%d segmentos)', len(plaintext), len(uncovered_segments))

    model_config = get_model_config()
    use_direct = model_config.get('use_direct_api', False)
    use_opencode = model_config.get('use_opencode', True)
    if not uncovered_segments:
        pii_keywords, reasoning_output, queue_notice, ai_status = [], '', None, 'skipped'
        logger.info('Todos los segmentos cubiertos por regex+AymurAI. IA omitida.')
    elif use_direct:
        pii_keywords, reasoning_output, queue_notice, ai_status = call_direct_api_for_pii(
            plaintext,
            wait_seconds=wait_seconds,
        )
    elif not use_opencode:
        pii_keywords, reasoning_output, queue_notice, ai_status = [], '', None, 'skipped'
        logger.info('OpenCode deshabilitado en config. IA omitida.')
    else:
        pii_keywords, reasoning_output, queue_notice, ai_status = call_opencode_for_pii(
            plaintext,
            wait_seconds=wait_seconds,
        )
    if pii_keywords:
        logger.info('IA devolvio %d keywords (status=%s)', len(pii_keywords), ai_status)
    else:
        logger.info('IA no devolvio keywords (status=%s)', ai_status)

    ai_positions = find_word_positions(segments, pii_keywords)
    logger.info('Posiciones IA: %d', len(ai_positions))

    seen_ranges = set()
    merged_positions = []
    for pos in default_positions + aymurai_positions + ai_positions:
        key = (pos['segment'], pos['start'], pos['end'], pos['word'])
        if key not in seen_ranges:
            seen_ranges.add(key)
            merged_positions.append(pos)

    merged_positions.sort(key=lambda p: (p['segment'], p['start']))
    logger.info('Total posiciones fusionadas: %d', len(merged_positions))

    return {
        'keywords': aymurai_keywords + pii_keywords,
        'default_keywords': default_keywords,
        'positions': merged_positions,
        'reasoning': reasoning_output,
        'queue_notice': queue_notice,
        'ai_status': ai_status,
        'analysis_mode': 'full' if ai_status in ('ok', 'skipped') else 'regex_only',
        'ai_positions': ai_positions,
        'aymurai_positions': aymurai_positions,
    }


def normalize_text(text):
    nfkd = unicodedata.normalize('NFKD', text)
    return ''.join(c for c in nfkd if not unicodedata.combining(c))


def detect_default_pii(segments):
    default_keywords_set = {}
    positions = []
    patterns = get_pii_patterns()

    for seg_idx, seg in enumerate(segments):
        text = seg['text']
        lower_text = text.lower()

        for pattern, ptype in patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                word = match.group()
                normalized = normalize_text(word).lower()
                key = f'{normalized}|{ptype}'
                if key not in default_keywords_set:
                    default_keywords_set[key] = {'word': word, 'type': ptype}
                positions.append({
                    'segment': seg_idx,
                    'start': match.start(),
                    'end': match.end(),
                    'word': word,
                    'type': ptype
                })

    return list(default_keywords_set.values()), positions


def find_word_positions(segments, keywords):
    positions = []
    for seg_idx, seg in enumerate(segments):
        text = seg['text']
        norm_text = normalize_text(text).lower()
        for kw in keywords:
            norm_kw = normalize_text(kw['word']).lower()
            start = 0
            while True:
                idx = norm_text.find(norm_kw, start)
                if idx == -1:
                    break
                word_match = text[idx:idx + len(kw['word'])]
                positions.append({
                    'segment': seg_idx,
                    'start': idx,
                    'end': idx + len(kw['word']),
                    'word': word_match,
                    'type': kw.get('type', 'other')
                })
                start = idx + 1
    return positions


def replace_normalized(text, kw_word, replacement):
    norm_text = normalize_text(text).lower()
    norm_kw = normalize_text(kw_word).lower()
    if not norm_kw:
        return text
    start = 0
    while True:
        idx = norm_text.find(norm_kw, start)
        if idx == -1:
            break
        orig_text = text[idx:idx + len(kw_word)]
        match_len = len(orig_text)
        text = text[:idx] + replacement + text[idx + match_len:]
        norm_text = normalize_text(text).lower()
        start = idx + len(replacement)
    return text


def build_normalized_index_map(text):
    normalized_chars = []
    index_map = []
    for idx, char in enumerate(text):
        normalized = normalize_text(char)
        if not normalized:
            continue
        for norm_char in normalized:
            normalized_chars.append(norm_char)
            index_map.append(idx)
    return ''.join(normalized_chars), index_map


def find_normalized_ranges(text, kw_word):
    norm_text, index_map = build_normalized_index_map(text)
    norm_text = norm_text.lower()
    norm_kw = normalize_text(kw_word).lower()
    if not norm_kw:
        return []

    ranges = []
    start = 0
    while True:
        idx = norm_text.find(norm_kw, start)
        if idx == -1:
            break
        end_idx = idx + len(norm_kw) - 1
        if end_idx < len(index_map):
            orig_start = index_map[idx]
            orig_end = index_map[end_idx] + 1
            ranges.append((orig_start, orig_end))
        start = idx + len(norm_kw)
    return ranges


def replace_keywords_in_paragraph_runs(paragraph, keywords, replacement):
    if not paragraph.runs:
        return False

    matches = []
    full_text = paragraph.text
    for kw in keywords:
        for start, end in find_normalized_ranges(full_text, kw['word']):
            matches.append((start, end))
    if not matches:
        return False

    deduped = []
    seen = set()
    for start, end in matches:
        if (start, end) not in seen:
            deduped.append((start, end))
            seen.add((start, end))

    for start, end in sorted(deduped, reverse=True):
        run_positions = []
        cursor = 0
        for run in paragraph.runs:
            run_len = len(run.text or '')
            run_positions.append((run, cursor, cursor + run_len))
            cursor += run_len

        start_info = None
        end_info = None
        for run, run_start, run_end in run_positions:
            if start_info is None and run_start <= start < run_end:
                start_info = (run, run_start, run_end)
            if run_start < end <= run_end:
                end_info = (run, run_start, run_end)
                break

        if not start_info or not end_info:
            continue

        start_run, start_run_start, _ = start_info
        end_run, end_run_start, _ = end_info
        start_offset = start - start_run_start
        end_offset = end - end_run_start

        if start_run is end_run:
            text = start_run.text or ''
            start_run.text = text[:start_offset] + replacement + text[end_offset:]
            continue

        start_run.text = (start_run.text or '')[:start_offset] + replacement
        end_run.text = (end_run.text or '')[end_offset:]

        clearing = False
        for run, _, _ in run_positions:
            if run is start_run:
                clearing = True
                continue
            if run is end_run:
                break
            if clearing:
                run.text = ''

    return True


def iter_docx_blocks(parent):
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def sanitize_pdf_text(text):
    text = (text or '').replace('\xa0', ' ')
    text = re.sub(r'\s+', ' ', text).strip()

    def split_long_token(match):
        token = match.group(0)
        return ' '.join(token[i:i + 24] for i in range(0, len(token), 24))

    return re.sub(r'\S{40,}', split_long_token, text)


def safe_pdf_multi_cell(pdf, line_height, text):
    safe_text = sanitize_pdf_text(text)
    if not safe_text:
        return
    available_width = max(pdf.w - pdf.l_margin - pdf.r_margin, 10)
    try:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(available_width, line_height, safe_text)
        return
    except FPDFException:
        pass

    # Fallback for pathological lines/tables that FPDF cannot wrap.
    for chunk in re.split(r'\s*\|\s*', safe_text):
        chunk = sanitize_pdf_text(chunk)
        if not chunk:
            continue
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(available_width, line_height, chunk)


def anonymize_docx(filepath, keywords, replacement='[REDACTADO]'):
    doc = Document(filepath)

    def replace_text(text):
        out = text
        for kw in keywords:
            if normalize_text(kw['word']).lower() in normalize_text(out).lower():
                out = replace_normalized(out, kw['word'], replacement)
        return out

    def keep_run_style(paragraph):
        if paragraph.runs:
            src = paragraph.runs[0]
            return {
                'name': src.font.name,
                'size': src.font.size,
                'bold': src.font.bold,
                'italic': src.font.italic,
                'underline': src.font.underline,
                'color': src.font.color.rgb,
            }
        return None

    def apply_run_style(run, style):
        if not style:
            return
        run.font.name = style['name']
        run.font.size = style['size']
        run.font.bold = style['bold']
        run.font.italic = style['italic']
        run.font.underline = style['underline']
        run.font.color.rgb = style['color']

    def anonymize_paragraphs(paragraphs):
        for para in paragraphs:
            original = para.text
            if not original:
                continue
            anonymized = replace_text(original)
            if anonymized == original:
                continue

            if replace_keywords_in_paragraph_runs(para, keywords, replacement) and para.text == anonymized:
                continue

            style = keep_run_style(para)
            para.text = anonymized
            if para.runs:
                apply_run_style(para.runs[0], style)

    anonymize_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                anonymize_paragraphs(cell.paragraphs)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def get_pdf_fonts():
    import glob
    candidates = glob.glob('/usr/share/fonts/truetype/dejavu/DejaVuSans*.ttf')
    regular = None
    bold = None
    for f in sorted(candidates):
        basename = f.split('/')[-1]
        if basename == 'DejaVuSans.ttf':
            regular = f
        elif basename == 'DejaVuSans-Bold.ttf':
            bold = f
    return regular, bold, None


def anonymize_pdf(segments, keywords, replacement='[REDACTADO]', title='Documento Anonimizado'):
    pdf = FPDF(format='A4')
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(18, 18, 18)
    pdf.add_page()

    regular_font, bold_font, _ = get_pdf_fonts()
    has_dejavu = regular_font is not None

    if has_dejavu:
        pdf.add_font('DejaVuSans', '', regular_font, uni=True)
        pdf.add_font('DejaVuSans', 'B', bold_font, uni=True)

    if has_dejavu:
        pdf.set_font('DejaVuSans', 'B', 15)
    else:
        pdf.set_font('Helvetica', 'B', 15)
    safe_pdf_multi_cell(pdf, 8, title)
    pdf.ln(2)

    for seg in segments:
        text = seg['text']
        for kw in keywords:
            if normalize_text(kw['word']).lower() in normalize_text(text).lower():
                text = replace_normalized(text, kw['word'], replacement)

        if has_dejavu:
            if seg['type'] == 'title':
                pdf.set_font('DejaVuSans', 'B', 13)
            elif seg['type'] == 'list':
                pdf.set_font('DejaVuSans', '', 11)
            else:
                pdf.set_font('DejaVuSans', '', 11)
        else:
            if seg['type'] == 'title':
                pdf.set_font('Helvetica', 'B', 13)
            elif seg['type'] == 'list':
                pdf.set_font('Helvetica', '', 11)
            else:
                pdf.set_font('Helvetica', '', 11)

        if seg['type'] == 'list':
            text = f'- {text}'

        line_h = 6.6 if seg['type'] == 'paragraph' else 7.2
        safe_pdf_multi_cell(pdf, line_h, text)
        if seg['type'] == 'title':
            pdf.ln(2.5)
        else:
            pdf.ln(1.4)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


def anonymize_docx_to_pdf(filepath, keywords, replacement='[REDACTADO]', title='Documento Anonimizado'):
    doc = Document(filepath)
    pdf = FPDF(format='A4')
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(18, 18, 18)
    pdf.add_page()

    regular_font, bold_font, _ = get_pdf_fonts()
    has_dejavu = regular_font is not None

    if has_dejavu:
        pdf.add_font('DejaVuSans', '', regular_font, uni=True)
        pdf.add_font('DejaVuSans', 'B', bold_font, uni=True)

    def set_font(style='', size=11):
        if has_dejavu:
            pdf.set_font('DejaVuSans', style, size)
        else:
            pdf.set_font('Helvetica', style, size)

    def replace_text(text):
        out = text
        for kw in keywords:
            if normalize_text(kw['word']).lower() in normalize_text(out).lower():
                out = replace_normalized(out, kw['word'], replacement)
        return out

    set_font('B', 15)
    safe_pdf_multi_cell(pdf, 8, title)
    pdf.ln(2)

    for block in iter_docx_blocks(doc):
        if isinstance(block, Paragraph):
            text = replace_text(block.text.strip())
            if not text:
                continue
            style_name = block.style.name.lower() if block.style else ''
            if 'heading 1' in style_name or 'title' in style_name:
                set_font('B', 15)
                safe_pdf_multi_cell(pdf, 8, text)
                pdf.ln(2)
            elif 'heading' in style_name:
                set_font('B', 13)
                safe_pdf_multi_cell(pdf, 7, text)
                pdf.ln(1.5)
            elif 'list' in style_name:
                set_font('', 11)
                safe_pdf_multi_cell(pdf, 6.6, f'- {text}')
                pdf.ln(1)
            else:
                is_bold = any(run.bold for run in block.runs)
                set_font('B' if is_bold else '', 11)
                safe_pdf_multi_cell(pdf, 6.6, text)
                pdf.ln(1.2)
        elif isinstance(block, Table):
            set_font('', 10)
            for row in block.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = ' '.join(
                        replace_text(p.text.strip()) for p in cell.paragraphs if p.text.strip()
                    )
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    safe_pdf_multi_cell(pdf, 6.2, ' | '.join(row_texts))
            pdf.ln(1.5)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/ready', methods=['GET'])
def ready():
    inflight = _current_inflight()
    busy = inflight >= READY_MAX_INFLIGHT
    payload = {
        'ready': not busy,
        'busy': busy,
        'inflight': inflight,
        'max_inflight': READY_MAX_INFLIGHT,
    }
    return (jsonify(payload), 200) if not busy else (jsonify(payload), 503)


@app.route('/upload', methods=['POST'])
def upload():
    logger.info('--- Nuevo upload ---')
    cleanup_old_uploads()
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    file = request.files['file']
    if not file.filename or not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type. Only PDF and DOCX allowed'}), 400

    ext = file.filename.rsplit('.', 1)[1].lower()
    uid = str(uuid.uuid4())
    filename = f'{uid}.{ext}'
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    logger.info('Archivo guardado: %s (%s)', filename, ext)

    segments, used_ocr = extract_text(filepath)
    if not segments:
        return jsonify({'error': 'Could not extract text from this document'}), 400
    logger.info('Texto extraído: %d segmentos', len(segments))

    result = run_detection_pipeline(segments)
    logger.info('--- Upload completado ---')

    return jsonify({
        'filename': filename,
        'segments': segments,
        'keywords': result['keywords'],
        'default_keywords': result['default_keywords'],
        'positions': result['positions'],
        'reasoning': result['reasoning'],
        'queue_notice': result['queue_notice'],
        'ai_status': result['ai_status'],
        'analysis_mode': result['analysis_mode'],
        'used_ocr': used_ocr,
    })


@app.route('/reanalyze-ai', methods=['POST'])
def reanalyze_ai():
    cleanup_old_uploads()
    data = request.get_json() or {}
    filename = data.get('filename', '')
    if not is_valid_upload_filename(filename):
        return jsonify({'error': 'Invalid filename'}), 400

    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not is_path_inside_uploads(filepath) or not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404

    segments, used_ocr = extract_text(filepath)
    if not segments:
        return jsonify({'error': 'Could not extract text from this document'}), 400

    result = run_detection_pipeline(segments, wait_seconds=0)
    status_code = 202 if result['ai_status'] in ('busy', 'unavailable') else 200
    return jsonify({
        'filename': filename,
        'keywords': result['keywords'],
        'default_keywords': result['default_keywords'],
        'positions': result['positions'],
        'ai_positions': result['ai_positions'],
        'reasoning': result['reasoning'],
        'queue_notice': result['queue_notice'],
        'ai_status': result['ai_status'],
        'analysis_mode': result['analysis_mode'],
        'used_ocr': used_ocr,
    }), status_code


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('admin_logged_in'):
            return jsonify({'error': 'Unauthorized'}), 401
        return f(*args, **kwargs)
    return decorated


@app.route('/uploads/<filename>')
@admin_required
def uploaded_file(filename):
    if not is_valid_upload_filename(filename):
        return jsonify({'error': 'Invalid filename'}), 400
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not is_path_inside_uploads(filepath) or not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)


@app.route('/export', methods=['POST'])
def export():
    logger.info('--- Export solicitado ---')
    cleanup_old_uploads()
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    filename = data.get('filename')
    keywords = data.get('keywords', [])
    output_format = data.get('format', 'docx')
    replacement = data.get('replacement', '[REDACTADO]')

    if not filename:
        return jsonify({'error': 'No filename provided'}), 400

    if not is_valid_upload_filename(filename):
        return jsonify({'error': 'Invalid filename'}), 400

    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not is_path_inside_uploads(filepath):
        return jsonify({'error': 'Invalid path'}), 400
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404

    ext = filename.rsplit('.', 1)[1].lower()

    if ext == 'pdf' and output_format == 'docx':
        return jsonify({'error': 'No se puede exportar PDF a DOCX. Usá formato PDF.'}), 400

    kw_entries = [{'word': kw['word'], 'type': kw.get('type', 'other')}
                  for kw in keywords]
    logger.info('Exportando %s con %d keywords a reemplazar (formato: %s)', filename, len(kw_entries), output_format)

    if output_format == 'docx':
        buf = anonymize_docx(filepath, kw_entries, replacement)
        logger.info('Export DOCX completado')
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument'
                      '.wordprocessingml.document',
            as_attachment=True,
            download_name=f'anonimizado_{filename.replace(".pdf", ".docx")}'
        )
    elif output_format == 'pdf':
        if ext == 'docx':
            buf = anonymize_docx_to_pdf(
                filepath,
                kw_entries,
                replacement,
                f'Anonimizado - {filename}'
            )
        else:
            segments, _ = extract_text(filepath)
            buf = anonymize_pdf(segments, kw_entries, replacement,
                               f'Anonimizado - {filename}')
        logger.info('Export PDF completado')
        return send_file(
            buf,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'anonimizado_{filename.replace(".docx", ".pdf")}'
        )

    return jsonify({'error': 'Unsupported format. Use docx or pdf'}), 400


@app.route('/admin/login', methods=['POST'])
def admin_login():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    ip = request.headers.get('X-Forwarded-For', request.remote_addr or 'unknown').split(',')[0].strip()
    if is_login_rate_limited(ip):
        return jsonify({'error': 'Too many attempts. Try again later.'}), 429
    user = data.get('user', '')
    password = data.get('password', '')
    if hmac.compare_digest(str(user), str(ADMIN_USER)) and hmac.compare_digest(str(password), str(ADMIN_PASS)):
        session['admin_logged_in'] = True
        if redis_available:
            try:
                redis_client.delete(f'anonimizador:login:{ip}')
            except redis.RedisError:
                pass
        logger.info('Admin login exitoso')
        return jsonify({'ok': True})
    register_login_failure(ip)
    logger.warning('Intento de login admin fallido')
    return jsonify({'error': 'Credenciales inválidas'}), 401


@app.route('/admin/logout', methods=['POST'])
def admin_logout():
    session.pop('admin_logged_in', None)
    session.modified = True
    return jsonify({'ok': True})


@app.route('/admin/status', methods=['GET'])
def admin_status():
    return jsonify({'logged_in': session.get('admin_logged_in', False)})


@app.route('/admin/config', methods=['GET'])
@admin_required
def admin_get_config():
    config = load_regex_config()
    return jsonify({
        'patterns': config.get('patterns', []),
        'prompt': config.get('prompt', ''),
        'model_url': config.get('model_url', os.environ.get('OPENAI_BASE_URL', '')),
        'model_name': config.get('model_name', os.environ.get('MODEL_NAME', '')),
        'api_key': config.get('api_key', os.environ.get('OPENAI_API_KEY', '')),
        'opencode_command': config.get(
            'opencode_command',
            'opencode run "{message}" --model opencode/{model} --dangerously-skip-permissions --file {file}'
        ),
        'use_direct_api': config.get('use_direct_api', False),
        'use_opencode': config.get('use_opencode', True),
        'use_aymurai': config.get('use_aymurai', USE_AYMURAI),
        'aymurai_url': config.get('aymurai_url', AYMURAI_BASE_URL),
    })


def validate_config_patterns(patterns):
    if not isinstance(patterns, list):
        return 'patterns debe ser una lista'
    for i, p in enumerate(patterns):
        if not isinstance(p, dict):
            return f'patrón {i}: debe ser un objeto JSON'
        if 'pattern' not in p or 'type' not in p:
            return f'patrón {i}: requiere campos "pattern" y "type"'
        if not isinstance(p['pattern'], str) or not isinstance(p['type'], str):
            return f'patrón {i}: "pattern" y "type" deben ser strings'
        if len(p['pattern']) > 300:
            return f'patrón {i}: regex demasiado largo (máx 300 chars)'
        try:
            re.compile(p['pattern'])
        except re.error as e:
            return f'patrón {i}: regex inválido — {e}'
    return None


def validate_config_prompt(prompt):
    if not isinstance(prompt, str):
        return 'prompt debe ser un string'
    if len(prompt) > 10000:
        return 'prompt demasiado largo (máx 10000 chars)'
    return None


def validate_config_model_url(model_url):
    if model_url is None:
        return None
    if not isinstance(model_url, str):
        return 'model_url debe ser un string'
    if len(model_url) > 500:
        return 'model_url demasiado largo (máx 500 chars)'
    if model_url.strip():
        try:
            from urllib.parse import urlparse
            parsed = urlparse(model_url)
            if parsed.scheme not in ('http', 'https'):
                return 'model_url debe usar http o https'
            if not parsed.netloc:
                return 'model_url debe incluir un host válido'
        except Exception:
            return 'model_url no es una URL válida'
    return None


def validate_opencode_command(opencode_command):
    if opencode_command is None:
        return None
    if not isinstance(opencode_command, str):
        return 'opencode_command debe ser un string'
    if len(opencode_command) > 1000:
        return 'opencode_command demasiado largo (máx 1000 chars)'
    text = opencode_command.strip()
    if not text:
        return None
    if '{file}' not in text:
        return 'opencode_command debe incluir placeholder {file}'
    return None


def validate_config_api_key(api_key):
    if api_key is None:
        return None
    if not isinstance(api_key, str):
        return 'api_key debe ser un string'
    if len(api_key) > 500:
        return 'api_key demasiado larga (máx 500 chars)'
    return None


def validate_config_use_direct_api(use_direct_api):
    if use_direct_api is None:
        return None
    return None


def validate_config_use_opencode(use_opencode):
    if use_opencode is None:
        return None
    return None


def validate_config_use_aymurai(use_aymurai):
    if use_aymurai is None:
        return None
    return None


def validate_config_aymurai_url(aymurai_url):
    if aymurai_url is None:
        return None
    if not isinstance(aymurai_url, str):
        return 'aymurai_url debe ser un string'
    if len(aymurai_url) > 500:
        return 'aymurai_url demasiado larga (máx 500 chars)'
    if aymurai_url.strip():
        from urllib.parse import urlparse
        parsed = urlparse(aymurai_url)
        if parsed.scheme not in ('http', 'https'):
            return 'aymurai_url debe usar http o https'
        if not parsed.netloc:
            return 'aymurai_url debe incluir un host válido'
    return None


@app.route('/admin/config', methods=['POST'])
@admin_required
def admin_save_config():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    patterns = data.get('patterns', [])
    prompt = data.get('prompt', '')
    model_url = data.get('model_url')
    model_name = data.get('model_name')
    api_key = data.get('api_key')
    opencode_command = data.get('opencode_command')
    use_direct_api = data.get('use_direct_api')
    use_opencode = data.get('use_opencode')
    use_aymurai = data.get('use_aymurai')
    aymurai_url = data.get('aymurai_url')

    err = validate_config_patterns(patterns)
    if err:
        return jsonify({'error': err}), 400
    err = validate_config_prompt(prompt)
    if err:
        return jsonify({'error': err}), 400
    err = validate_config_model_url(model_url)
    if err:
        return jsonify({'error': err}), 400
    err = validate_opencode_command(opencode_command)
    if err:
        return jsonify({'error': err}), 400
    err = validate_config_api_key(api_key)
    if err:
        return jsonify({'error': err}), 400
    err = validate_config_use_direct_api(use_direct_api)
    if err:
        return jsonify({'error': err}), 400
    err = validate_config_use_opencode(use_opencode)
    if err:
        return jsonify({'error': err}), 400
    err = validate_config_use_aymurai(use_aymurai)
    if err:
        return jsonify({'error': err}), 400
    err = validate_config_aymurai_url(aymurai_url)
    if err:
        return jsonify({'error': err}), 400

    try:
        save_regex_config(patterns, prompt, model_url, model_name, opencode_command, api_key, use_direct_api, use_opencode, use_aymurai, aymurai_url)
        return jsonify({'ok': True})
    except Exception as e:
        logger.error('Error guardando config: %s', e)
        return jsonify({'error': str(e)}), 500


@app.route('/admin/aymurai-status', methods=['GET'])
@admin_required
def admin_aymurai_status():
    config = load_regex_config()
    aym_url = get_aymurai_url()
    use_aymurai_config = config.get('use_aymurai', USE_AYMURAI)
    result = {
        'enabled': bool(aym_url),
        'use_aymurai': use_aymurai_config,
        'url': aym_url,
        'available': False,
        'error': None,
    }
    if not result['enabled']:
        return jsonify(result)
    api_url = f'{aym_url}/anonymizer/predict'
    try:
        req = Request(api_url, data=b'{"text":"test"}', headers={'Content-Type': 'application/json'}, method='POST')
        with urlopen(req, timeout=AYMURAI_TIMEOUT_SECONDS) as resp:
            if resp.status == 200:
                result['available'] = True
    except Exception as e:
        result['error'] = str(e)
    return jsonify(result)


@app.route('/admin/test-api', methods=['POST'])
@admin_required
def admin_test_api():
    payload = request.get_json(silent=True) or {}
    cfg = get_model_config()
    base_url = (payload.get('model_url') if 'model_url' in payload else cfg.get('model_url') or '').strip().rstrip('/')
    if not base_url:
        return jsonify({'ok': False, 'error': 'No hay model_url configurado'}), 400
    use_direct_api = payload.get('use_direct_api') if 'use_direct_api' in payload else cfg.get('use_direct_api')
    if not use_direct_api:
        return jsonify({'ok': False, 'error': 'El modo API directa no esta habilitado'}), 400

    model_name = payload.get('model_name') if 'model_name' in payload else cfg.get('model_name', '')
    api_key = payload.get('api_key') if 'api_key' in payload else cfg.get('api_key', '')
    api_key = (api_key or '').strip()

    api_url = f"{base_url}/chat/completions"
    payload = json.dumps({
        "model": model_name,
        "messages": [{"role": "user", "content": "Hi"}],
        "temperature": 0,
        "max_tokens": 5,
    }).encode('utf-8')

    headers = {'Content-Type': 'application/json'}
    if api_key:
        headers['Authorization'] = f'Bearer {api_key}'

    started = time.time()
    try:
        req = Request(api_url, data=payload, headers=headers, method='POST')
        with urlopen(req, timeout=15) as resp:
            duration_ms = int((time.time() - started) * 1000)
            body = resp.read().decode('utf-8', errors='replace')
            data_json = json.loads(body)
            content = data_json.get('choices', [{}])[0].get('message', {}).get('content', '')
            _log_api_call(base_url, model_name, 'test-ok', duration_ms, content[:200])
            return jsonify({
                'ok': True,
                'message': f'Conexion exitosa ({duration_ms}ms)',
                'response': content[:200],
            })
    except Exception as e:
        duration_ms = int((time.time() - started) * 1000)
        _log_api_call(base_url, model_name, f'test-error: {e}', duration_ms)
        return jsonify({
            'ok': False,
            'error': str(e),
        }), 502


@app.route('/admin/test-inference', methods=['POST'])
@admin_required
def admin_test_inference():
    prompt_text = 'hola, explica tu capacidades'
    content, error, status = run_model_inference(prompt_text)
    if error:
        return jsonify({
            'ok': False,
            'error': error,
            'status': status,
        }), 502
    return jsonify({
        'ok': True,
        'message': 'Inferencia completada',
        'response': content,
        'status': status,
    })


@app.route('/admin/api-logs', methods=['GET'])
@admin_required
def admin_api_logs():
    logs = _read_api_logs_from_redis() if redis_available else []
    if not logs:
        logs = list(_api_call_log[-50:])
    return jsonify({'logs': logs[-50:]})


@app.route('/admin/config/restore-defaults', methods=['POST'])
@admin_required
def admin_restore_defaults():
    current = load_regex_config()
    save_regex_config(
        patterns=DEFAULT_PATTERNS_DATA['patterns'],
        prompt=DEFAULT_PATTERNS_DATA['prompt'],
        model_url=current.get('model_url'),
        model_name=current.get('model_name'),
        api_key=current.get('api_key'),
        opencode_command=current.get('opencode_command'),
        use_direct_api=current.get('use_direct_api', False),
    )
    return jsonify({'ok': True})


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)
