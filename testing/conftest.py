import os
import sys
import json
import re
import tempfile
import uuid
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

os.environ['FLASK_SECRET_KEY'] = 'test-secret-key'
os.environ['ADMIN_USER'] = 'admin'
os.environ['ADMIN_PASS'] = 'testpass'
os.environ['MODEL_NAME'] = 'opencode/test-model'
os.environ['OPENAI_BASE_URL'] = 'https://api.test.com/v1'
os.environ['FLASK_PORT'] = '5000'
os.environ['SESSION_COOKIE_SECURE'] = '0'
import sys
print(f'[CONFTEST] Setting REDIS_URL override', file=sys.stderr)
os.environ['REDIS_URL'] = 'redis://redis:6379/15'
_TEST_REDIS_CONFIG_KEY = f"anonimizador:config:test:{uuid.uuid4()}"
os.environ['REDIS_CONFIG_KEY'] = _TEST_REDIS_CONFIG_KEY

print(f'[CONFTEST] REDIS_URL={os.environ["REDIS_URL"]}, REDIS_CONFIG_KEY={os.environ["REDIS_CONFIG_KEY"]}', file=sys.stderr)
sys.stderr.flush()

import app as anon_app

REGEX_PATTERNS_FILE = anon_app.REGEX_PATTERNS_FILE
DEFAULT_PATTERNS_DATA = anon_app.DEFAULT_PATTERNS_DATA

def reset_config_file():
    with open(REGEX_PATTERNS_FILE, 'w', encoding='utf-8') as f:
        json.dump(DEFAULT_PATTERNS_DATA, f, indent=2, ensure_ascii=False)
    if anon_app.redis_available:
        try:
            anon_app.redis_client.delete(anon_app.REDIS_CONFIG_KEY)
        except Exception:
            pass

reset_config_file()

def create_synthetic_docx(paragraphs, tables=None, filepath=None):
    import docx
    from docx import Document
    doc = Document()
    for para in paragraphs:
        if isinstance(para, dict):
            if para.get('style') == 'heading':
                doc.add_heading(para['text'], level=para.get('level', 1))
            elif para.get('style') == 'list':
                doc.add_paragraph(para['text'], style='List Bullet')
            else:
                doc.add_paragraph(para['text'])
        else:
            doc.add_paragraph(para)
    if tables:
        for table_data in tables:
            rows = len(table_data)
            cols = max(len(r) for r in table_data)
            table = doc.add_table(rows=rows, cols=cols)
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    table.cell(i, j).text = cell_text
    if filepath:
        doc.save(filepath)
        return filepath
    buf = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(buf.name)
    buf.close()
    return buf.name

def create_synthetic_pdf(text_content, filepath=None):
    from fpdf import FPDF
    pdf = FPDF(format='A4')
    pdf.add_page()
    pdf.set_font('Helvetica', '', 11)
    for line in text_content.split('\n'):
        line = line.strip()
        if line:
            pdf.multi_cell(0, 6, line)
    if filepath:
        pdf.output(filepath)
        return filepath
    buf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
    pdf.output(buf.name)
    buf.close()
    return buf.name

def pii_doc_paragraphs():
    return [
        {'text': 'Pericia Psicológica', 'style': 'heading', 'level': 1},
        {'text': 'Paciente: Juan Carlos Martínez'},
        {'text': 'DNI 30.123.456'},
        {'text': 'Domicilio: Calle San Martín 1234, CABA'},
        {'text': 'Edad: 35 años'},
        {'text': 'Sexo: Masculino'},
        {'text': 'Expediente N° 12345/2024'},
        {'text': 'La víctima, María López, denunció violencia de género.'},
        {'text': 'El imputado, Roberto Gómez, fue detenido el 15/01/2024.'},
        {'text': 'Se trata de un hecho de abuso sexual contra una menor de edad.'},
        {'text': 'La pericia forense determinó lesiones compatibles con el relato.'},
        {'text': 'El fallecimiento ocurrió por homicidio agravado.'},
        {'text': 'El Juzgado de Familia N° 3 interviene en la causa.'},
        {'text': 'Email de contacto: juan.perez@email.com'},
        {'text': 'CUIL 20-30123456-7'},
    ]

def pii_doc_esperado():
    return {
        'dni': ['30.123.456'],
        'nombres': ['Juan Carlos Martínez', 'María López', 'Roberto Gómez'],
        'direccion': ['Calle San Martín 1234'],
        'edad': ['35 años'],
        'sexo': ['Masculino'],
        'expedientes': ['12345/2024'],
        'email': ['juan.perez@email.com'],
    }

import pytest

@pytest.fixture(autouse=True)
def _reset_config_for_each_test():
    reset_config_file()
    yield

def cleanup_temp_files(*paths):
    for p in paths:
        try:
            os.unlink(p)
        except (OSError, TypeError):
            pass
