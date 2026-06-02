import pdfplumber
from io import BytesIO
from conftest import anon_app

def _extract_pdf_text(buf):
    text = ''
    with pdfplumber.open(buf) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ''
            text += page_text + '\n'
    return text

def test_anonymize_pdf_replaces_keywords():
    segments = [{'type': 'paragraph', 'text': 'Paciente: Juan Pérez, DNI 30.123.456'}]
    keywords = [
        {'word': 'Juan Pérez', 'type': 'nombre'},
        {'word': '30.123.456', 'type': 'dni'},
    ]
    buf = anon_app.anonymize_pdf(segments, keywords, '[REDACTADO]')
    assert isinstance(buf, BytesIO)
    assert len(buf.getvalue()) > 0
    content = _extract_pdf_text(buf)
    assert 'Juan Pérez' not in content
    assert '30.123.456' not in content
    assert 'Paciente:' in content

def test_anonymize_pdf_preserves_non_pii():
    segments = [
        {'type': 'paragraph', 'text': 'Informe médico general.'},
        {'type': 'paragraph', 'text': 'El paciente se encuentra en buen estado.'},
    ]
    keywords = [{'word': 'Juan Pérez', 'type': 'nombre'}]
    buf = anon_app.anonymize_pdf(segments, keywords, '[REDACTADO]')
    content = _extract_pdf_text(buf)
    assert 'Informe médico general' in content
    assert 'buen estado' in content

def test_anonymize_pdf_empty_keywords():
    segments = [{'type': 'paragraph', 'text': 'Texto de prueba sin datos.'}]
    buf = anon_app.anonymize_pdf(segments, [], '[REDACTADO]')
    content = _extract_pdf_text(buf)
    assert 'Texto de prueba sin datos.' in content

def test_anonymize_pdf_title_segment():
    segments = [
        {'type': 'title', 'text': 'Pericia Psicológica - Juan Pérez'},
        {'type': 'paragraph', 'text': 'Paciente: Juan Pérez'},
    ]
    keywords = [{'word': 'Juan Pérez', 'type': 'nombre'}]
    buf = anon_app.anonymize_pdf(segments, keywords, '[REDACTADO]')
    content = _extract_pdf_text(buf)
    assert 'Juan Pérez' not in content
    assert 'Pericia Psicológica' in content

def test_anonymize_pdf_list_segment():
    segments = [
        {'type': 'list', 'text': 'Juan Pérez - 30.123.456'},
        {'type': 'list', 'text': 'María López - 40.123.456'},
    ]
    keywords = [
        {'word': 'Juan Pérez', 'type': 'nombre'},
        {'word': '30.123.456', 'type': 'dni'},
        {'word': 'María López', 'type': 'nombre'},
        {'word': '40.123.456', 'type': 'dni'},
    ]
    buf = anon_app.anonymize_pdf(segments, keywords, '[REDACTADO]')
    content = _extract_pdf_text(buf)
    for kw in ['Juan Pérez', '30.123.456', 'María López', '40.123.456']:
        assert kw not in content
    assert '- ' in content

def test_anonymize_pdf_multiple_segments():
    segments = [
        {'type': 'paragraph', 'text': 'Informe de Pericia'},
        {'type': 'paragraph', 'text': 'Paciente: Pedro Gómez'},
        {'type': 'paragraph', 'text': 'DNI: 40.123.456'},
    ]
    keywords = [
        {'word': 'Pedro Gómez', 'type': 'nombre'},
        {'word': '40.123.456', 'type': 'dni'},
    ]
    buf = anon_app.anonymize_pdf(segments, keywords, '[REDACTADO]')
    content = _extract_pdf_text(buf)
    assert 'Pedro Gómez' not in content
    assert '40.123.456' not in content
    assert 'Informe de Pericia' in content

def test_anonymize_pdf_custom_replacement():
    segments = [{'type': 'paragraph', 'text': 'Paciente: Juan Pérez'}]
    keywords = [{'word': 'Juan Pérez', 'type': 'nombre'}]
    buf = anon_app.anonymize_pdf(segments, keywords, '[ANONIMO]')
    content = _extract_pdf_text(buf)
    assert '[ANONIMO]' in content
    assert 'Juan Pérez' not in content

def test_anonymize_pdf_custom_title():
    segments = [{'type': 'paragraph', 'text': 'Paciente: Juan Pérez'}]
    keywords = [{'word': 'Juan Pérez', 'type': 'nombre'}]
    buf = anon_app.anonymize_pdf(segments, keywords, '[REDACTADO]', 'Documento Anonimizado - Test')
    content = _extract_pdf_text(buf)
    assert 'Documento Anonimizado - Test' in content

def test_anonymize_pdf_accented():
    segments = [{'type': 'paragraph', 'text': 'El Dr. José Martínez realizó la pericia.'}]
    keywords = [{'word': 'José Martínez', 'type': 'nombre'}]
    buf = anon_app.anonymize_pdf(segments, keywords, '[REDACTADO]')
    content = _extract_pdf_text(buf)
    assert 'José Martínez' not in content
    assert 'El Dr.' in content

def test_anonymize_pdf_output_is_bytesio():
    segments = [{'type': 'paragraph', 'text': 'Paciente: Juan Pérez'}]
    keywords = [{'word': 'Juan Pérez', 'type': 'nombre'}]
    buf = anon_app.anonymize_pdf(segments, keywords, '[REDACTADO]')
    assert isinstance(buf, BytesIO)
    assert len(buf.getvalue()) > 0

def test_anonymize_pdf_scanned_pdf_ocr_fallback():
    import os
    scan_path = os.path.join(os.path.dirname(__file__), 'scansmpl.pdf')
    assert os.path.exists(scan_path)
    
    segments, used_ocr = anon_app.extract_text_pdf(scan_path)
    assert used_ocr is True
    assert len(segments) > 0
    
    keywords = []
    for seg in segments:
        text = seg['text']
        if len(text) > 10:
            words = text.split()
            if len(words) >= 2:
                keywords.append({'word': ' '.join(words[:2]), 'type': 'test'})
    
    if keywords:
        buf = anon_app.anonymize_pdf(segments, keywords, '[REDACTADO]')
        assert isinstance(buf, BytesIO)
        assert len(buf.getvalue()) > 0
        content = _extract_pdf_text(buf)
        assert content

def test_anonymize_docx_to_pdf_preserves_heading_and_table_content():
    import os
    import tempfile
    from docx import Document

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = Document()
        doc.add_heading('Informe Clínico', level=1)
        doc.add_paragraph('Paciente: Juan Pérez')
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = 'DNI'
        table.cell(0, 1).text = '30.123.456'
        doc.save(tmp.name)

        keywords = [
            {'word': 'Juan Pérez', 'type': 'nombre'},
            {'word': '30.123.456', 'type': 'dni'},
        ]
        buf = anon_app.anonymize_docx_to_pdf(tmp.name, keywords, '[REDACTADO]', 'Anonimizado - Test')
        content = _extract_pdf_text(buf)

        assert 'Informe Clínico' in content
        assert 'Juan Pérez' not in content
        assert '30.123.456' not in content
        assert 'DNI' in content
    finally:
        os.unlink(tmp.name)


def test_anonymize_pdf_smart_replacement_preserves_prefix_and_consistency():
    segments = [
        {'type': 'paragraph', 'text': 'Paciente: Juan Pérez'},
        {'type': 'paragraph', 'text': 'Paciente: Juan Pérez'},
    ]
    keywords = [{'word': 'Paciente: Juan Pérez', 'type': 'nombre'}]
    buf = anon_app.anonymize_pdf(segments, keywords, '[REDACTADO]')
    content = _extract_pdf_text(buf)
    assert 'Paciente:' in content
    assert 'Juan Pérez' not in content
    assert '[REDACTADO]' not in content
    assert content.count('Paciente:') >= 2


def test_anonymize_pdf_unsupported_type_falls_back_to_redactado():
    segments = [{'type': 'paragraph', 'text': 'Expediente N° 12345/2024'}]
    keywords = [{'word': 'Expediente N° 12345/2024', 'type': 'sensible'}]
    buf = anon_app.anonymize_pdf(segments, keywords, '[REDACTADO]')
    content = _extract_pdf_text(buf)
    assert '[REDACTADO]' in content
    assert 'Expediente N° 12345/2024' not in content
