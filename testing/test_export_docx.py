import os
import tempfile
from conftest import anon_app

def test_anonymize_docx_replaces_keywords():
    import docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        doc.add_paragraph('Paciente: Juan Pérez, DNI 30.123.456')
        doc.save(tmp.name)
        keywords = [
            {'word': 'Juan Pérez', 'type': 'nombre'},
            {'word': '30.123.456', 'type': 'dni'},
        ]
        buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
        result_doc = docx.Document(buf)
        text = ' '.join(p.text for p in result_doc.paragraphs)
        assert 'Juan Pérez' not in text
        assert '30.123.456' not in text
        assert text != 'Paciente: Juan Pérez, DNI 30.123.456'
    finally:
        os.unlink(tmp.name)

def test_anonymize_docx_preserves_non_pii():
    import docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        doc.add_paragraph('Informe médico general.')
        doc.add_paragraph('El paciente se encuentra en buen estado.')
        doc.save(tmp.name)
        keywords = [{'word': 'Juan Pérez', 'type': 'nombre'}]
        buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
        result_doc = docx.Document(buf)
        texts = [p.text for p in result_doc.paragraphs]
        assert any('Informe médico general' in t for t in texts)
        assert any('buen estado' in t for t in texts)
    finally:
        os.unlink(tmp.name)

def test_anonymize_docx_empty_keywords():
    import docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        original = 'Texto de prueba sin datos sensibles.'
        doc.add_paragraph(original)
        doc.save(tmp.name)
        buf = anon_app.anonymize_docx(tmp.name, [], '[REDACTADO]')
        result_doc = docx.Document(buf)
        assert original in ' '.join(p.text for p in result_doc.paragraphs)
    finally:
        os.unlink(tmp.name)

def test_anonymize_docx_multiple_paragraphs():
    import docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        doc.add_paragraph('Paciente: Pedro Gómez')
        doc.add_paragraph('DNI: 40.123.456')
        doc.add_paragraph('Diagnóstico: nada relevante.')
        doc.save(tmp.name)
        keywords = [
            {'word': 'Pedro Gómez', 'type': 'nombre'},
            {'word': '40.123.456', 'type': 'dni'},
        ]
        buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
        result_doc = docx.Document(buf)
        full = ' '.join(p.text for p in result_doc.paragraphs)
        assert 'Pedro Gómez' not in full
        assert '40.123.456' not in full
        assert 'nada relevante' in full
    finally:
        os.unlink(tmp.name)

def test_anonymize_docx_table_cells():
    import docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'Nombre'
        table.cell(0, 1).text = 'DNI'
        table.cell(1, 0).text = 'Juan Pérez'
        table.cell(1, 1).text = '30.123.456'
        doc.save(tmp.name)
        keywords = [
            {'word': 'Juan Pérez', 'type': 'nombre'},
            {'word': '30.123.456', 'type': 'dni'},
        ]
        buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
        result_doc = docx.Document(buf)
        full = ' '.join(p.text for p in result_doc.paragraphs)
        for table in result_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full += ' ' + cell.text
        assert 'Juan Pérez' not in full
        assert '30.123.456' not in full
        assert full != ' Nombre DNI Juan Pérez 30.123.456'
    finally:
        os.unlink(tmp.name)

def test_anonymize_docx_replacement_string():
    import docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        doc.add_paragraph('Paciente: Juan Pérez')
        doc.save(tmp.name)
        keywords = [{'word': 'Juan Pérez', 'type': 'nombre'}]
        buf = anon_app.anonymize_docx(tmp.name, keywords, '[ANONIMO]')
        result_doc = docx.Document(buf)
        text = ' '.join(p.text for p in result_doc.paragraphs)
        assert '[ANONIMO]' in text
        assert '[REDACTADO]' not in text
    finally:
        os.unlink(tmp.name)

def test_anonymize_docx_accented_keyword():
    import docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        doc.add_paragraph('El perito es José Martínez')
        doc.save(tmp.name)
        keywords = [{'word': 'José Martínez', 'type': 'nombre'}]
        buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
        result_doc = docx.Document(buf)
        text = ' '.join(p.text for p in result_doc.paragraphs)
        assert 'José Martínez' not in text
        assert text != 'El perito es José Martínez'
    finally:
        os.unlink(tmp.name)

def test_anonymize_docx_output_is_bytesio():
    import docx
    from io import BytesIO
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        doc.add_paragraph('Paciente: Juan Pérez')
        doc.save(tmp.name)
        keywords = [{'word': 'Juan Pérez', 'type': 'nombre'}]
        buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
        assert isinstance(buf, BytesIO)
        assert len(buf.getvalue()) > 0
    finally:
        os.unlink(tmp.name)

def test_anonymize_docx_preserves_other_run_formatting():
    import docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        para = doc.add_paragraph()
        run1 = para.add_run('Paciente: Juan Pérez')
        run1.bold = True
        run2 = para.add_run(' - Observación importante')
        run2.italic = True
        doc.save(tmp.name)

        keywords = [{'word': 'Juan Pérez', 'type': 'nombre'}]
        buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
        result_doc = docx.Document(buf)
        result_para = result_doc.paragraphs[0]

        assert 'Juan Pérez' not in result_para.text
        assert result_para.runs[0].bold is True
        assert any(run.italic for run in result_para.runs if 'Observación importante' in run.text)
    finally:
        os.unlink(tmp.name)

def test_anonymize_docx_preserves_format_when_keyword_spans_runs():
    import docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        para = doc.add_paragraph()
        run1 = para.add_run('Juan ')
        run1.bold = True
        run2 = para.add_run('Pérez')
        run2.bold = True
        run3 = para.add_run(' - Observación importante')
        run3.italic = True
        doc.save(tmp.name)

        keywords = [{'word': 'Juan Pérez', 'type': 'nombre'}]
        buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
        result_doc = docx.Document(buf)
        result_para = result_doc.paragraphs[0]

        assert 'Juan Pérez' not in result_para.text
        assert any(run.bold for run in result_para.runs if run.text.strip())
        assert any(run.italic for run in result_para.runs if 'Observación importante' in run.text)
    finally:
        os.unlink(tmp.name)


def test_anonymize_docx_smart_replacement_preserves_prefix_and_consistency():
    import docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        doc.add_paragraph('Paciente: Juan Pérez')
        doc.add_paragraph('Paciente: Juan Pérez')
        doc.save(tmp.name)
        keywords = [{'word': 'Paciente: Juan Pérez', 'type': 'nombre'}]
        buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
        result_doc = docx.Document(buf)
        texts = [p.text for p in result_doc.paragraphs]
        assert texts[0].startswith('Paciente: ')
        assert texts[0] == texts[1]
        assert 'Juan Pérez' not in texts[0]
        assert '[REDACTADO]' not in texts[0]
    finally:
        os.unlink(tmp.name)


def test_anonymize_docx_unsupported_type_falls_back_to_redactado():
    import docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        doc.add_paragraph('Expediente N° 12345/2024')
        doc.save(tmp.name)
        keywords = [{'word': 'Expediente N° 12345/2024', 'type': 'sensible'}]
        buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
        result_doc = docx.Document(buf)
        text = ' '.join(p.text for p in result_doc.paragraphs)
        assert '[REDACTADO]' in text
        assert 'Expediente N° 12345/2024' not in text
    finally:
        os.unlink(tmp.name)


def test_build_position_replacements_is_stable_for_same_word():
    positions = [
        {'segment': 0, 'start': 0, 'end': 11, 'word': 'Juan Pérez', 'type': 'nombre'},
        {'segment': 1, 'start': 0, 'end': 11, 'word': 'Juan Pérez', 'type': 'nombre'},
    ]
    replacements = anon_app.build_position_replacements(positions)
    assert len(replacements) == 2
    assert replacements[0]['replacement'] == replacements[1]['replacement']
    assert replacements[0]['replacement'] != '[REDACTADO]'


def test_build_position_replacements_falls_back_for_unsupported_type():
    positions = [
        {'segment': 0, 'start': 0, 'end': 24, 'word': 'Expediente N° 12345/2024', 'type': 'sensible'},
    ]
    replacements = anon_app.build_position_replacements(positions)
    assert replacements == [{
        'segment': 0,
        'start': 0,
        'end': 24,
        'replacement': '[REDACTADO]',
    }]


def test_build_position_replacements_supports_alias_types():
    positions = [
        {'segment': 0, 'start': 0, 'end': 13, 'word': 'Claudio Perez', 'type': 'persona'},
        {'segment': 0, 'start': 20, 'end': 39, 'word': 'Av. Siempre Viva 742', 'type': 'domicilio'},
    ]
    replacements = anon_app.build_position_replacements(positions)
    assert replacements[0]['replacement'] != '[REDACTADO]'
    assert replacements[1]['replacement'] != '[REDACTADO]'


def test_build_position_replacements_supports_judicial_entities():
    positions = [
        {'segment': 0, 'start': 0, 'end': 20, 'word': 'Juzgado de Familia', 'type': 'juzgado'},
        {'segment': 0, 'start': 21, 'end': 42, 'word': 'Fiscalía Federal N° 2', 'type': 'fiscalia'},
        {'segment': 0, 'start': 43, 'end': 64, 'word': 'Defensoría Oficial N° 1', 'type': 'defensoria'},
        {'segment': 0, 'start': 65, 'end': 73, 'word': 'T. 54 F. 233', 'type': 'matricula_prof'},
        {'segment': 0, 'start': 74, 'end': 88, 'word': '011-4567-8910', 'type': 'telefono'},
    ]
    replacements = anon_app.build_position_replacements(positions)
    assert all(item['replacement'] != '[REDACTADO]' for item in replacements)


def test_build_position_replacements_supports_dates_and_case_numbers():
    positions = [
        {'segment': 0, 'start': 0, 'end': 10, 'word': '15/03/2024', 'type': 'fecha'},
        {'segment': 0, 'start': 11, 'end': 25, 'word': 'EXP-2024-55432', 'type': 'expediente_judicial'},
        {'segment': 0, 'start': 26, 'end': 44, 'word': 'CPACF T. 54 F. 233', 'type': 'matricula_prof'},
        {'segment': 0, 'start': 45, 'end': 60, 'word': '+54 11 4567-8901', 'type': 'telefono'},
    ]
    replacements = anon_app.build_position_replacements(positions)
    assert all(item['replacement'] != '[REDACTADO]' for item in replacements)


def test_anonymize_docx_preserves_judicial_role_prefix_in_name():
    import docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    try:
        doc = docx.Document()
        doc.add_paragraph('Juez Claudio Pérez intervino en la audiencia.')
        doc.save(tmp.name)
        keywords = [{'word': 'Juez Claudio Pérez', 'type': 'nombre'}]
        buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
        result_doc = docx.Document(buf)
        text = ' '.join(p.text for p in result_doc.paragraphs)
        assert 'Claudio Pérez' not in text
        assert text.startswith('Juez ')
        assert '[REDACTADO]' not in text
    finally:
        os.unlink(tmp.name)
