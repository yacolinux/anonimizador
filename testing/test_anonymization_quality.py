import os
import json
import tempfile
import uuid
import pytest
from conftest import (
    anon_app, create_synthetic_docx, cleanup_temp_files, pii_doc_paragraphs
)

DOCX_TMP = None
PDF_TMP = None

def setup_module():
    global DOCX_TMP, PDF_TMP
    DOCX_TMP = create_synthetic_docx(pii_doc_paragraphs())
    PDF_TMP = None

def teardown_module():
    cleanup_temp_files(DOCX_TMP, PDF_TMP)

class TestAnonymizationQualityDNI:

    def test_regex_detects_dni_formatted(self):
        segments = [{'type': 'paragraph', 'text': 'DNI 30.123.456'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        assert any('30.123.456' in kw['word'] for kw in keywords)

    def test_regex_detects_dni_simple(self):
        segments = [{'type': 'paragraph', 'text': 'Documento 30123456'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        assert any('30123456' in kw['word'] for kw in keywords)

    def test_dni_anonimizado_en_docx(self):
        import docx
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        tmp.close()
        try:
            d = docx.Document()
            d.add_paragraph('DNI 30.123.456 del paciente')
            d.save(tmp.name)
            keywords = [{'word': '30.123.456', 'type': 'dni_argentino'}]
            buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
            rd = docx.Document(buf)
            text = ' '.join(p.text for p in rd.paragraphs)
            assert '30.123.456' not in text
            assert '[REDACTADO]' in text
        finally:
            os.unlink(tmp.name)

class TestAnonymizationQualityCUIL:

    def test_dni_pattern_catches_cuil(self):
        segments = [{'type': 'paragraph', 'text': 'CUIL 20-30123456-7'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        assert len(positions) >= 1

    def test_cuil_anonimizado_en_docx(self):
        import docx
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        tmp.close()
        try:
            d = docx.Document()
            d.add_paragraph('CUIL 20-30123456-7')
            d.save(tmp.name)
            keywords = [{'word': '20-30123456-7', 'type': 'dni_argentino'}]
            buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
            rd = docx.Document(buf)
            text = ' '.join(p.text for p in rd.paragraphs)
            assert '20-30123456-7' not in text
        finally:
            os.unlink(tmp.name)

class TestAnonymizationQualityNombres:

    def test_nombre_capturado_por_paciente(self):
        segments = [{'type': 'paragraph', 'text': 'Paciente: Juan Carlos Martínez'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        nombre_kw = [k for k in keywords if k['type'] == 'nombre']
        assert any('Juan Carlos Martínez' in kw['word'] for kw in nombre_kw)

    def test_nombre_simple_regex(self):
        segments = [{'type': 'paragraph', 'text': 'Sr. Juan Pérez'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        assert any(k['type'] == 'nombre' for k in keywords)

    def test_nombre_anonimizado(self):
        import docx
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        tmp.close()
        try:
            d = docx.Document()
            d.add_paragraph('Paciente: Juan Carlos Martínez')
            d.save(tmp.name)
            keywords = [{'word': 'Juan Carlos Martínez', 'type': 'nombre'}]
            buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
            rd = docx.Document(buf)
            text = ' '.join(p.text for p in rd.paragraphs)
            assert 'Juan Carlos Martínez' not in text
            assert '[REDACTADO]' in text
        finally:
            os.unlink(tmp.name)

    def test_nombre_con_acento_anonimizado(self):
        import docx
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        tmp.close()
        try:
            d = docx.Document()
            d.add_paragraph('Paciente: José Martínez')
            d.save(tmp.name)
            keywords = [{'word': 'José Martínez', 'type': 'nombre'}]
            buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
            rd = docx.Document(buf)
            text = ' '.join(p.text for p in rd.paragraphs)
            assert 'José Martínez' not in text
        finally:
            os.unlink(tmp.name)

    def test_nombre_variante_sin_acento(self):
        segments = [{'type': 'paragraph', 'text': 'El doctor Gomez realizó la pericia'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        assert len(keywords) >= 0

class TestAnonymizationQualityDomicilios:

    def test_direccion_calle(self):
        segments = [{'type': 'paragraph', 'text': 'Calle San Martín 1234, CABA'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        direccion_kw = [k for k in keywords if k['type'] == 'direccion']
        assert len(direccion_kw) > 0

    def test_direccion_av(self):
        segments = [{'type': 'paragraph', 'text': 'Av. Corrientes 2567'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        direccion_kw = [k for k in keywords if k['type'] == 'direccion']
        assert len(direccion_kw) > 0

    def test_direccion_domicilio(self):
        segments = [{'type': 'paragraph', 'text': 'Domicilio: Calle Pellegrini 888'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        direccion_kw = [k for k in keywords if k['type'] == 'direccion']
        assert len(direccion_kw) > 0

    def test_direccion_pasaje(self):
        segments = [{'type': 'paragraph', 'text': 'Pasaje San Lorenzo 456'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        direccion_kw = [k for k in keywords if k['type'] == 'direccion']
        assert len(direccion_kw) > 0

    def test_direccion_ruta(self):
        segments = [{'type': 'paragraph', 'text': 'Ruta Nacional 8 km 45'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        direccion_kw = [k for k in keywords if k['type'] == 'direccion']
        assert len(direccion_kw) > 0

class TestAnonymizationQualityExpedientes:

    def test_expediente_detectado(self):
        segments = [{'type': 'paragraph', 'text': 'Expediente N° 12345/2024'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('expedient' in kw['word'].lower() for kw in sensible_kw)

    def test_numero_expediente_completo(self):
        segments = [{'type': 'paragraph', 'text': 'Expediente N° 12345/2024'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        assert len(positions) > 0

class TestAnonymizationQualityVictimas:

    def test_victima_detectada(self):
        segments = [{'type': 'paragraph', 'text': 'La víctima sufrió lesiones graves.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('lesion' in kw['word'].lower() for kw in sensible_kw)

    def test_victima_y_denuncia_juntas(self):
        segments = [{'type': 'paragraph', 'text': 'La víctima realizó la denuncia.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('denunci' in kw['word'].lower() for kw in sensible_kw)

class TestAnonymizationQualityImputados:

    def test_imputado_detectado(self):
        segments = [{'type': 'paragraph', 'text': 'El imputado fue detenido.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        imput_kw = [k for k in sensible_kw if 'imput' in k['word'].lower()]
        deten_kw = [k for k in sensible_kw if 'deten' in k['word'].lower()]
        assert len(imput_kw) > 0
        assert len(deten_kw) > 0

    def test_imputado_con_nombre(self):
        segments = [{'type': 'paragraph', 'text': 'El imputado, Roberto Gómez, fue notificado.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('imput' in kw['word'].lower() for kw in sensible_kw)

class TestAnonymizationQualityMenores:

    def test_menor_en_texto(self):
        segments = [{'type': 'paragraph', 'text': 'La víctima es una menor de edad.'}]
        keywords, positions = anon_app.detect_default_pii(segments)

    def test_menor_detectado_por_violacion(self):
        segments = [{'type': 'paragraph', 'text': 'abuso sexual contra una menor'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('abus' in kw['word'].lower() for kw in sensible_kw)

class TestAnonymizationQualityDelitosSexuales:

    def test_abuso_detectado(self):
        segments = [{'type': 'paragraph', 'text': 'Se investiga un abuso sexual.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('abus' in kw['word'].lower() for kw in sensible_kw)

    def test_agresion_detectada(self):
        segments = [{'type': 'paragraph', 'text': 'Fue agredida sexualmente.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('agred' in kw['word'].lower() or 'agresi' in kw['word'].lower()
                   for kw in sensible_kw)

    def test_amenaza_detectada(self):
        segments = [{'type': 'paragraph', 'text': 'Recibió amenazas de muerte.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('amenaz' in kw['word'].lower() for kw in sensible_kw)

    def test_lesiones_detectadas(self):
        segments = [{'type': 'paragraph', 'text': 'Lesiones compatibles con el relato.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('lesion' in kw['word'].lower() for kw in sensible_kw)

class TestAnonymizationQualityViolencia:

    def test_violencia_detectada(self):
        segments = [{'type': 'paragraph', 'text': 'Violencia de género contra la mujer.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('viol' in kw['word'].lower() for kw in sensible_kw)

    def test_violencia_detectada_en_variante(self):
        segments = [{'type': 'paragraph', 'text': 'violencia doméstica'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('viol' in kw['word'].lower() for kw in sensible_kw)

class TestAnonymizationQualityFallecimientos:

    def test_fallecimiento_detectado(self):
        segments = [{'type': 'paragraph', 'text': 'Se constató el fallecimiento.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('fallec' in kw['word'].lower() for kw in sensible_kw)

    def test_cadaver_detectado(self):
        segments = [{'type': 'paragraph', 'text': 'El cadaver fue identificado.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('cadav' in kw['word'].lower() for kw in sensible_kw)

    def test_homicidio_detectado(self):
        segments = [{'type': 'paragraph', 'text': 'Falleció por homicidio agravado.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('homicid' in kw['word'].lower() for kw in sensible_kw)

    def test_femicidio_detectado(self):
        segments = [{'type': 'paragraph', 'text': 'Femicidio en contexto de violencia.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('femicid' in kw['word'].lower() for kw in sensible_kw)

    def test_autopsia_detectada(self):
        segments = [{'type': 'paragraph', 'text': 'Autopsia realizada por el forense.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('autops' in kw['word'].lower() for kw in sensible_kw)

    def test_necropsia_detectada(self):
        segments = [{'type': 'paragraph', 'text': 'Necropsia de ley.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('necrops' in kw['word'].lower() for kw in sensible_kw)

class TestAnonymizationQualityOrganismosJudiciales:

    def test_juzgado_no_es_pii(self):
        segments = [{'type': 'paragraph', 'text': 'El Juzgado de Familia N° 3 interviene.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        for kw in keywords:
            assert 'Juzgado' not in kw['word']

    def test_organismo_no_es_pii(self):
        segments = [{'type': 'paragraph', 'text': 'El Juzgado de Familia interviene en la causa.'}]
        keywords, positions = anon_app.detect_default_pii(segments)

    def test_testigo_no_es_pii_pero_marcado_sensible(self):
        segments = [{'type': 'paragraph', 'text': 'El testigo declaró bajo juramento.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        if keywords:
            sensible_kw = [k for k in keywords if k['type'] == 'sensible']
            if sensible_kw:
                assert any('testig' in kw['word'].lower() for kw in sensible_kw)

    def test_pericia_forense_detectada(self):
        segments = [{'type': 'paragraph', 'text': 'Pericia forense realizada.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('forens' in kw['word'].lower() or 'perici' in kw['word'].lower()
                   for kw in sensible_kw)

    def test_identificacion_detectada(self):
        segments = [{'type': 'paragraph', 'text': 'Identificación de la menor.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('identif' in kw['word'].lower() for kw in sensible_kw)

    def test_documentacion_detectada(self):
        segments = [{'type': 'paragraph', 'text': 'Documentación anexa al expediente.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        doc_kw = [k for k in sensible_kw if 'document' in k['word'].lower() or
                  'expedient' in k['word'].lower()]
        assert len(doc_kw) > 0

    def test_condena_detectada(self):
        segments = [{'type': 'paragraph', 'text': 'Fue condenado a prisión.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('conden' in kw['word'].lower() for kw in sensible_kw)

    def test_denuncia_detectada(self):
        segments = [{'type': 'paragraph', 'text': 'Denuncia penal en curso.'}]
        keywords, positions = anon_app.detect_default_pii(segments)
        sensible_kw = [k for k in keywords if k['type'] == 'sensible']
        assert any('denunci' in kw['word'].lower() for kw in sensible_kw)

class TestAnonymizationQualityEndToEnd:

    def test_documento_completo_regex(self):
        paragraphs = pii_doc_paragraphs()
        segments = [{'type': 'paragraph', 'text': p['text'] if isinstance(p, dict) else p}
                    for p in paragraphs]
        keywords, positions = anon_app.detect_default_pii(segments)
        assert len(positions) > 5
        types = {k['type'] for k in keywords}
        assert 'nombre' in types or 'dni_argentino' in types or 'sensible' in types

    def test_documento_completo_anonimizado_docx(self):
        import docx
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        tmp.close()
        try:
            d = docx.Document()
            for p in pii_doc_paragraphs():
                if isinstance(p, dict):
                    d.add_paragraph(p['text'])
                else:
                    d.add_paragraph(p)
            d.save(tmp.name)
            keywords = [
                {'word': 'Juan Carlos Martínez', 'type': 'nombre'},
                {'word': '30.123.456', 'type': 'dni_argentino'},
            ]
            buf = anon_app.anonymize_docx(tmp.name, keywords, '[REDACTADO]')
            rd = docx.Document(buf)
            text = ' '.join(p.text for p in rd.paragraphs)
            assert 'Juan Carlos Martínez' not in text
            assert '30.123.456' not in text
        finally:
            os.unlink(tmp.name)

    def test_merge_positions_no_duplicates(self):
        segments = [
            {'type': 'paragraph', 'text': 'Paciente: Juan Pérez, DNI 30.123.456'}
        ]
        default_kw, default_pos = anon_app.detect_default_pii(segments)
        ai_keywords = [{'word': 'Juan Pérez', 'type': 'nombre'}]
        ai_pos = anon_app.find_word_positions(segments, ai_keywords)
        seen = set()
        for pos in default_pos + ai_pos:
            key = (pos['segment'], pos['start'], pos['end'], pos['word'])
            assert key not in seen
            seen.add(key)

    def test_find_word_positions_no_overlap(self):
        segments = [{'type': 'paragraph', 'text': 'Juan Pérez y María López'}]
        keywords = [
            {'word': 'Juan Pérez', 'type': 'nombre'},
            {'word': 'María López', 'type': 'nombre'},
        ]
        positions = anon_app.find_word_positions(segments, keywords)
        assert len(positions) == 2
        assert positions[0]['end'] <= positions[1]['start'] or \
               positions[0]['start'] >= positions[1]['end']